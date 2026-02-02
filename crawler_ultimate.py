# -*- coding: utf-8 -*-
"""
小红书爬虫终极版 v5.0
功能：视频下载、评论爬取、正文内容、标签提取、博主爬取、数据可视化、Cookie管理
优化：性能提升、稳定性增强、UI改进
"""

import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, filedialog
import threading
import queue
import json
import os
import time
import random
import re
import zipfile
import sqlite3
from typing import Optional, List, Dict, Any, Tuple, Callable
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import quote
from datetime import datetime
from collections import Counter
from dataclasses import dataclass, field

import pandas as pd
import requests
from DrissionPage import ChromiumPage, ChromiumOptions

# 版本信息
VERSION = "5.1"
APP_NAME = f"小红书爬虫终极版 v{VERSION}"

# 可选依赖
try:
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')  # 非交互式后端
    HAS_MATPLOTLIB = True
except:
    HAS_MATPLOTLIB = False

try:
    from wordcloud import WordCloud
    import jieba
    HAS_WORDCLOUD = True
except:
    HAS_WORDCLOUD = False

try:
    from docx import Document
    from docx.shared import Inches
    HAS_DOCX = True
except:
    HAS_DOCX = False


@dataclass
class CrawlerConfig:
    """爬虫配置（使用dataclass提升可维护性）"""
    # 基础配置
    keyword: str = ""
    scroll_times: int = 10
    max_notes: int = 30
    parallel_downloads: int = 10
    retry_times: int = 2
    save_interval: int = 10
    
    # 爬取内容选项（默认全部开启）
    download_images: bool = True
    download_videos: bool = True
    get_all_images: bool = True
    get_content: bool = True
    get_tags: bool = True
    get_publish_time: bool = True
    get_comments: bool = True
    comments_count: int = 20
    get_interactions: bool = True
    
    # 爬取模式
    crawl_mode: str = "standard"  # standard/fast/turbo
    crawl_type: str = "keyword"   # keyword/blogger/hot
    blogger_url: str = ""
    
    # 筛选条件
    min_likes: int = 0
    max_likes: int = 999999
    note_type_filter: str = "全部"
    date_filter: str = "全部"
    
    # 导出选项
    export_format: str = "xlsx"
    export_to_db: bool = True
    db_path: str = "data/redbook.db"
    
    # 速度控制（元组默认值需要用field）
    click_delay: Tuple[float, float] = field(default_factory=lambda: (0.2, 0.4))
    scroll_delay: Tuple[float, float] = field(default_factory=lambda: (0.3, 0.5))
    
    # Cookie和日志
    save_cookies: bool = True
    cookies_file: str = "data/cookies.json"
    log_to_file: bool = True
    log_file: str = "data/crawler.log"
    
    # 配置文件路径
    config_file: str = "data/settings.json"
    
    def save_to_file(self):
        """保存配置到文件"""
        import json
        try:
            # 确保data目录存在
            os.makedirs("data", exist_ok=True)
            config_dict = {
                'keyword': self.keyword,
                'scroll_times': self.scroll_times,
                'max_notes': self.max_notes,
                'parallel_downloads': self.parallel_downloads,
                'retry_times': self.retry_times,
                'download_images': self.download_images,
                'download_videos': self.download_videos,
                'get_all_images': self.get_all_images,
                'get_content': self.get_content,
                'get_tags': self.get_tags,
                'get_publish_time': self.get_publish_time,
                'get_comments': self.get_comments,
                'comments_count': self.comments_count,
                'get_interactions': self.get_interactions,
                'crawl_mode': self.crawl_mode,
                'crawl_type': self.crawl_type,
                'blogger_url': self.blogger_url,
                'min_likes': self.min_likes,
                'max_likes': self.max_likes,
                'note_type_filter': self.note_type_filter,
                'date_filter': self.date_filter,
                'export_format': self.export_format,
                'export_to_db': self.export_to_db,
            }
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(config_dict, f, ensure_ascii=False, indent=2)
            print(f"[配置] 已保存到 {self.config_file}")
        except Exception as e:
            print(f"[配置] 保存失败: {e}")
    
    def load_from_file(self):
        """从文件加载配置"""
        import json
        if not os.path.exists(self.config_file):
            print(f"[配置] 配置文件不存在，使用默认设置")
            return False
        try:
            with open(self.config_file, 'r', encoding='utf-8') as f:
                config_dict = json.load(f)
            # 更新配置
            for key, value in config_dict.items():
                if hasattr(self, key):
                    setattr(self, key, value)
            print(f"[配置] 已加载上次设置 (max_notes={self.max_notes}, keyword={self.keyword})")
            return True
        except Exception as e:
            print(f"[配置] 加载失败: {e}")
            return False


class FileLogger:
    """文件日志记录器（线程安全）"""
    
    def __init__(self, log_file: str):
        self.log_file = log_file
        self._lock = threading.Lock()
        self._ensure_dir()
    
    def _ensure_dir(self):
        """确保日志目录存在"""
        log_dir = os.path.dirname(self.log_file)
        if log_dir:
            os.makedirs(log_dir, exist_ok=True)
        
    def log(self, message: str, level: str = "INFO"):
        """线程安全的日志写入"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        log_line = f"[{timestamp}] [{level}] {message}\n"
        with self._lock:
            try:
                with open(self.log_file, 'a', encoding='utf-8') as f:
                    f.write(log_line)
            except Exception:
                pass


class CookieManager:
    """Cookie管理器（支持过期检测）"""
    
    def __init__(self, cookies_file: str):
        self.cookies_file = cookies_file
        self._lock = threading.Lock()
    
    def _ensure_dir(self):
        """确保目录存在"""
        cookie_dir = os.path.dirname(self.cookies_file)
        if cookie_dir:
            os.makedirs(cookie_dir, exist_ok=True)
        
    def save(self, page) -> bool:
        """保存Cookie"""
        with self._lock:
            try:
                cookies = page.cookies()
                self._ensure_dir()
                # 添加保存时间戳
                data = {
                    'cookies': cookies,
                    'saved_at': datetime.now().isoformat(),
                    'version': VERSION
                }
                with open(self.cookies_file, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                return True
            except Exception:
                return False
    
    def load(self, page) -> bool:
        """加载Cookie"""
        with self._lock:
            try:
                if not os.path.exists(self.cookies_file):
                    return False
                    
                with open(self.cookies_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # 兼容旧格式
                cookies = data.get('cookies', data) if isinstance(data, dict) else data
                
                loaded = 0
                for cookie in cookies:
                    try:
                        page.set.cookies(cookie)
                        loaded += 1
                    except Exception:
                        pass
                return loaded > 0
            except Exception:
                return False
    
    def exists(self) -> bool:
        """检查Cookie是否存在"""
        return os.path.exists(self.cookies_file)
    
    def get_saved_time(self) -> Optional[str]:
        """获取Cookie保存时间"""
        try:
            if not os.path.exists(self.cookies_file):
                return None
            with open(self.cookies_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return data.get('saved_at', '未知')
        except Exception:
            return None
    
    def clear(self):
        """清除Cookie"""
        if os.path.exists(self.cookies_file):
            os.remove(self.cookies_file)


class DatabaseManager:
    """数据库管理器"""
    def __init__(self, db_path):
        self.db_path = db_path
        self._init_db()
    
    def _init_db(self):
        """初始化数据库"""
        os.makedirs(os.path.dirname(self.db_path), exist_ok=True)
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS notes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                note_id TEXT UNIQUE,
                title TEXT,
                author TEXT,
                content TEXT,
                tags TEXT,
                publish_time TEXT,
                ip_region TEXT,
                like_count INTEGER,
                collect_count INTEGER,
                comment_count INTEGER,
                note_type TEXT,
                note_link TEXT,
                image_urls TEXT,
                video_url TEXT,
                comments TEXT,
                keyword TEXT,
                crawl_time TEXT
            )
        ''')
        
        conn.commit()
        conn.close()
    
    def insert_note(self, note_data):
        """插入笔记"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            cursor.execute('''
                INSERT OR REPLACE INTO notes 
                (note_id, title, author, content, tags, publish_time, ip_region,
                 like_count, collect_count, comment_count, note_type, note_link,
                 image_urls, video_url, comments, keyword, crawl_time)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                note_data.get('note_id', ''),
                note_data.get('title', ''),
                note_data.get('author', ''),
                note_data.get('content', ''),
                json.dumps(note_data.get('tags', []), ensure_ascii=False),
                note_data.get('publish_time', ''),
                note_data.get('ip_region', ''),
                note_data.get('like_count', 0),
                note_data.get('collect_count', 0),
                note_data.get('comment_count', 0),
                note_data.get('note_type', ''),
                note_data.get('note_link', ''),
                json.dumps(note_data.get('image_urls', []), ensure_ascii=False),
                note_data.get('video_url', ''),
                json.dumps(note_data.get('comments', []), ensure_ascii=False),
                note_data.get('keyword', ''),
                datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            ))
            conn.commit()
            return True
        except Exception as e:
            return False
        finally:
            conn.close()
    
    def get_existing_note_ids(self, keyword):
        """获取已存在的笔记ID（用于增量更新）"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('SELECT note_id FROM notes WHERE keyword = ?', (keyword,))
        ids = set(row[0] for row in cursor.fetchall())
        conn.close()
        return ids


class MediaDownloader:
    """高性能媒体下载器（支持图片和视频）"""
    
    # 常用User-Agent列表，随机选择以避免被封
    USER_AGENTS = [
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36',
        'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
    ]
    
    def __init__(self, max_workers: int = 10, retry_times: int = 2, timeout: int = 15):
        self.max_workers = max_workers
        self.retry_times = retry_times
        self.timeout = timeout
        self._session = None
        self._stats = {'success': 0, 'failed': 0, 'bytes': 0}
    
    @property
    def session(self) -> requests.Session:
        """懒加载Session，复用连接"""
        if self._session is None:
            self._session = requests.Session()
            self._session.headers.update({
                'User-Agent': random.choice(self.USER_AGENTS),
                'Referer': 'https://www.xiaohongshu.com/',
                'Accept': 'image/webp,image/apng,image/*,video/*,*/*;q=0.8',
                'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
                'Origin': 'https://www.xiaohongshu.com',
            })
        return self._session
    
    def set_cookies(self, cookies):
        """设置Cookie（用于需要认证的下载）"""
        if cookies:
            for cookie in cookies:
                self.session.cookies.set(
                    cookie.get('name', ''),
                    cookie.get('value', ''),
                    domain=cookie.get('domain', '.xiaohongshu.com')
                )
    
    def _normalize_url(self, url: str) -> str:
        """标准化URL"""
        if not url:
            return ""
        if url.startswith('//'):
            return 'https:' + url
        if not url.startswith('http'):
            return 'https://' + url
        return url
    
    def download_file(self, url: str, local_path: str, 
                      stop_flag: Optional[Callable] = None,
                      min_size: int = 1024) -> Optional[str]:
        """下载单个文件"""
        url = self._normalize_url(url)
        if not url:
            return None
            
        for attempt in range(self.retry_times):
            if stop_flag and stop_flag():
                return None
            try:
                response = self.session.get(url, timeout=self.timeout, stream=True)
                response.raise_for_status()
                
                # 确保目录存在
                os.makedirs(os.path.dirname(local_path), exist_ok=True)
                
                # 流式写入
                total_size = 0
                with open(local_path, 'wb') as f:
                    for chunk in response.iter_content(chunk_size=16384):
                        if stop_flag and stop_flag():
                            f.close()
                            if os.path.exists(local_path):
                                os.remove(local_path)
                            return None
                        if chunk:
                            f.write(chunk)
                            total_size += len(chunk)
                
                # 检查文件大小
                if total_size < min_size:
                    os.remove(local_path)
                    return None
                
                self._stats['success'] += 1
                self._stats['bytes'] += total_size
                return local_path
                
            except requests.Timeout:
                if attempt < self.retry_times - 1:
                    time.sleep(0.2 * (attempt + 1))
            except Exception:
                if attempt < self.retry_times - 1:
                    time.sleep(0.1)
        
        self._stats['failed'] += 1
        return None
    
    def download_batch(self, tasks: List[Tuple[str, str]], 
                       progress_callback: Optional[Callable] = None,
                       stop_flag: Optional[Callable] = None) -> Dict[str, Optional[str]]:
        """批量并行下载"""
        if not tasks:
            return {}
            
        results = {}
        completed = 0
        total = len(tasks)
        
        if stop_flag and stop_flag():
            return results
        
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            future_to_task = {}
            for url, path in tasks:
                if stop_flag and stop_flag():
                    break
                future = executor.submit(self.download_file, url, path, stop_flag)
                future_to_task[future] = (url, path)
            
            for future in as_completed(future_to_task):
                if stop_flag and stop_flag():
                    # 取消剩余任务
                    for f in future_to_task:
                        f.cancel()
                    break
                    
                url, path = future_to_task[future]
                try:
                    results[url] = future.result(timeout=self.timeout + 5)
                except Exception:
                    results[url] = None
                    
                completed += 1
                if progress_callback:
                    progress_callback(completed, total)
        
        return results
    
    def get_stats(self) -> Dict[str, int]:
        """获取下载统计"""
        return self._stats.copy()
    
    def reset_stats(self):
        """重置统计"""
        self._stats = {'success': 0, 'failed': 0, 'bytes': 0}
    
    def close(self):
        """关闭Session"""
        if self._session:
            self._session.close()
            self._session = None


class DataAnalyzer:
    """数据分析器"""
    
    @staticmethod
    def generate_stats(df):
        """生成统计数据"""
        stats = {
            'total_notes': len(df),
            'total_likes': df['like_count'].sum() if 'like_count' in df.columns else 0,
            'avg_likes': df['like_count'].mean() if 'like_count' in df.columns else 0,
            'max_likes': df['like_count'].max() if 'like_count' in df.columns else 0,
            'total_collects': df['collect_count'].sum() if 'collect_count' in df.columns else 0,
            'total_comments': df['comment_count'].sum() if 'comment_count' in df.columns else 0,
            'image_notes': len(df[df['note_type'] == '图文']) if 'note_type' in df.columns else 0,
            'video_notes': len(df[df['note_type'] == '视频']) if 'note_type' in df.columns else 0,
        }
        return stats
    
    @staticmethod
    def generate_charts(df, output_dir):
        """生成图表"""
        if not HAS_MATPLOTLIB:
            return []
        
        charts = []
        os.makedirs(output_dir, exist_ok=True)
        
        plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei']
        plt.rcParams['axes.unicode_minus'] = False
        
        try:
            # 点赞分布图
            if 'like_count' in df.columns:
                fig, ax = plt.subplots(figsize=(10, 6))
                df['like_count'].hist(bins=20, ax=ax, color='#ff6b6b', edgecolor='white')
                ax.set_title('点赞数分布', fontsize=14)
                ax.set_xlabel('点赞数')
                ax.set_ylabel('笔记数量')
                chart_path = os.path.join(output_dir, 'likes_distribution.png')
                plt.savefig(chart_path, dpi=100, bbox_inches='tight')
                plt.close()
                charts.append(chart_path)
            
            # 笔记类型饼图
            if 'note_type' in df.columns:
                fig, ax = plt.subplots(figsize=(8, 8))
                type_counts = df['note_type'].value_counts()
                ax.pie(type_counts.values, labels=type_counts.index, autopct='%1.1f%%',
                       colors=['#4ecdc4', '#ff6b6b', '#ffe66d'])
                ax.set_title('笔记类型分布', fontsize=14)
                chart_path = os.path.join(output_dir, 'type_distribution.png')
                plt.savefig(chart_path, dpi=100, bbox_inches='tight')
                plt.close()
                charts.append(chart_path)
            
            # Top10点赞笔记
            if 'like_count' in df.columns and 'title' in df.columns:
                fig, ax = plt.subplots(figsize=(12, 6))
                top10 = df.nlargest(10, 'like_count')
                titles = [t[:15] + '...' if len(t) > 15 else t for t in top10['title']]
                ax.barh(range(len(top10)), top10['like_count'], color='#667eea')
                ax.set_yticks(range(len(top10)))
                ax.set_yticklabels(titles)
                ax.set_xlabel('点赞数')
                ax.set_title('Top10 热门笔记', fontsize=14)
                ax.invert_yaxis()
                chart_path = os.path.join(output_dir, 'top10_notes.png')
                plt.savefig(chart_path, dpi=100, bbox_inches='tight')
                plt.close()
                charts.append(chart_path)
                
        except Exception as e:
            pass
        
        return charts
    
    @staticmethod
    def generate_wordcloud(texts, output_path):
        """生成词云"""
        if not HAS_WORDCLOUD:
            return None
        
        try:
            # 合并文本并分词
            all_text = ' '.join(texts)
            words = jieba.cut(all_text)
            word_list = [w for w in words if len(w) > 1]
            word_freq = Counter(word_list)
            
            # 生成词云
            wc = WordCloud(
                font_path='C:/Windows/Fonts/simhei.ttf',
                width=800,
                height=400,
                background_color='white',
                max_words=100,
                colormap='viridis'
            )
            wc.generate_from_frequencies(word_freq)
            
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            wc.to_file(output_path)
            return output_path
        except:
            return None
    
    @staticmethod
    def generate_report(df, stats, charts, output_path, keyword):
        """生成Word分析报告"""
        if not HAS_DOCX:
            return None
        
        try:
            doc = Document()
            doc.add_heading(f'小红书数据分析报告 - {keyword}', 0)
            doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            
            # 统计概览
            doc.add_heading('数据概览', level=1)
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'
            
            stats_items = [
                ('总笔记数', stats.get('total_notes', 0)),
                ('总点赞数', stats.get('total_likes', 0)),
                ('平均点赞', f"{stats.get('avg_likes', 0):.1f}"),
                ('最高点赞', stats.get('max_likes', 0)),
            ]
            
            for i, (label, value) in enumerate(stats_items):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            
            # 图表
            if charts:
                doc.add_heading('数据可视化', level=1)
                for chart in charts:
                    if os.path.exists(chart):
                        doc.add_picture(chart, width=Inches(6))
                        doc.add_paragraph('')
            
            # Top10列表
            doc.add_heading('热门笔记 Top10', level=1)
            if 'like_count' in df.columns:
                top10 = df.nlargest(10, 'like_count')
                for i, row in top10.iterrows():
                    title = row.get('title', '')[:50]
                    likes = row.get('like_count', 0)
                    doc.add_paragraph(f"• {title}... (点赞 {likes})")
            
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            return output_path
        except:
            return None


class CrawlerApp:
    """爬虫GUI应用"""
    
    def __init__(self):
        self.root = tk.Tk()
        self.root.title(APP_NAME)
        self.root.geometry("980x850")
        self.root.minsize(800, 600)
        
        self.config = CrawlerConfig()
        # 加载上次的配置
        self.config.load_from_file()
        
        self.downloader = MediaDownloader()
        self.cookie_mgr = CookieManager(self.config.cookies_file)
        self.file_logger = FileLogger(self.config.log_file)
        self.db_mgr = DatabaseManager(self.config.db_path)
        
        self.log_queue = queue.Queue()
        self.is_running = False
        self.should_stop = False
        self.all_notes_data = []
        self.current_crawl_dir = ""  # 当前爬取的目录
        self.batch_notes_data = []  # 批次笔记数据
        self.current_batch_folder = None  # 当前批次文件夹
        self.browser_page = None  # 保持浏览器实例，避免每次都重新登录
        
        self._create_ui()
        self._start_log_consumer()
        
        # 恢复上次的GUI设置
        self._restore_gui_settings()
        
        # 程序退出时关闭浏览器并保存配置
        self.root.protocol("WM_DELETE_WINDOW", self._on_closing)
    
    def _create_ui(self):
        """创建界面"""
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        self.notebook = notebook  # 保存引用
        
        # 创建各标签页
        main_page = ttk.Frame(notebook, padding="10")
        result_page = ttk.Frame(notebook, padding="10")
        content_page = ttk.Frame(notebook, padding="10")
        analysis_page = ttk.Frame(notebook, padding="10")
        settings_page = ttk.Frame(notebook, padding="10")
        
        notebook.add(main_page, text="搜索爬取")
        notebook.add(result_page, text="爬取结果")
        notebook.add(content_page, text="内容选项")
        notebook.add(analysis_page, text="数据分析")
        notebook.add(settings_page, text="高级设置")
        
        self._create_main_page(main_page)
        self._create_result_page(result_page)
        self._create_content_page(content_page)
        self._create_analysis_page(analysis_page)
        self._create_settings_page(settings_page)
    
    def _create_main_page(self, parent):
        """创建主页面"""
        # === 爬取模式选择 ===
        mode_frame = ttk.LabelFrame(parent, text="爬取模式", padding="10")
        mode_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.crawl_type_var = tk.StringVar(value="keyword")
        
        mode_row = ttk.Frame(mode_frame)
        mode_row.pack(fill=tk.X)
        
        ttk.Radiobutton(mode_row, text="关键词搜索", variable=self.crawl_type_var, 
                       value="keyword", command=self._on_mode_change).pack(side=tk.LEFT, padx=(0, 15))
        ttk.Radiobutton(mode_row, text="博主主页", variable=self.crawl_type_var, 
                       value="blogger", command=self._on_mode_change).pack(side=tk.LEFT, padx=(0, 15))
        ttk.Radiobutton(mode_row, text="热门榜单", variable=self.crawl_type_var, 
                       value="hot", command=self._on_mode_change).pack(side=tk.LEFT)
        
        # === 搜索配置 ===
        self.search_frame = ttk.LabelFrame(parent, text="搜索配置", padding="10")
        self.search_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 关键词输入
        row1 = ttk.Frame(self.search_frame)
        row1.pack(fill=tk.X, pady=2)
        
        ttk.Label(row1, text="搜索关键词:").pack(side=tk.LEFT)
        self.keyword_var = tk.StringVar(value="鞋子")
        self.keyword_entry = ttk.Entry(row1, textvariable=self.keyword_var, width=40)
        self.keyword_entry.pack(side=tk.LEFT, padx=5)
        
        ttk.Label(row1, text="(多个用逗号分隔)", foreground="gray").pack(side=tk.LEFT)
        
        # 博主URL输入
        row1b = ttk.Frame(self.search_frame)
        row1b.pack(fill=tk.X, pady=2)
        
        ttk.Label(row1b, text="博主主页URL:").pack(side=tk.LEFT)
        self.blogger_url_var = tk.StringVar()
        self.blogger_entry = ttk.Entry(row1b, textvariable=self.blogger_url_var, width=50)
        self.blogger_entry.pack(side=tk.LEFT, padx=5)
        self.blogger_entry.config(state=tk.DISABLED)
        
        # 热门分类
        row1c = ttk.Frame(self.search_frame)
        row1c.pack(fill=tk.X, pady=2)
        
        ttk.Label(row1c, text="热门分类:").pack(side=tk.LEFT)
        self.hot_category_var = tk.StringVar(value="综合")
        self.hot_combo = ttk.Combobox(row1c, textvariable=self.hot_category_var,
                                      values=["综合", "美食", "穿搭", "美妆", "旅行", "家居", "数码"], 
                                      width=15, state="readonly")
        self.hot_combo.pack(side=tk.LEFT, padx=5)
        self.hot_combo.config(state=tk.DISABLED)
        
        # 数量配置
        row2 = ttk.Frame(self.search_frame)
        row2.pack(fill=tk.X, pady=5)
        
        # 滚动次数已改为自动模式，无需手动设置
        self.scroll_var = tk.StringVar(value="10")  # 保留变量但不显示
        
        ttk.Label(row2, text="最多笔记:").pack(side=tk.LEFT)
        self.max_notes_var = tk.StringVar(value="30")
        ttk.Spinbox(row2, from_=1, to=500, textvariable=self.max_notes_var, width=6).pack(side=tk.LEFT, padx=(2, 15))
        
        ttk.Label(row2, text="并行下载:").pack(side=tk.LEFT)
        self.parallel_var = tk.StringVar(value="10")
        ttk.Spinbox(row2, from_=1, to=20, textvariable=self.parallel_var, width=6).pack(side=tk.LEFT)
        
        # === 筛选条件 ===
        filter_frame = ttk.LabelFrame(parent, text="筛选条件", padding="10")
        filter_frame.pack(fill=tk.X, pady=(0, 10))
        
        filter_row = ttk.Frame(filter_frame)
        filter_row.pack(fill=tk.X)
        
        ttk.Label(filter_row, text="点赞范围:").pack(side=tk.LEFT)
        self.min_likes_var = tk.StringVar(value="0")
        ttk.Entry(filter_row, textvariable=self.min_likes_var, width=8).pack(side=tk.LEFT, padx=2)
        ttk.Label(filter_row, text="-").pack(side=tk.LEFT)
        self.max_likes_var = tk.StringVar(value="999999")
        ttk.Entry(filter_row, textvariable=self.max_likes_var, width=8).pack(side=tk.LEFT, padx=(2, 15))
        
        ttk.Label(filter_row, text="笔记类型:").pack(side=tk.LEFT)
        self.note_type_var = tk.StringVar(value="全部")
        ttk.Combobox(filter_row, textvariable=self.note_type_var,
                    values=["全部", "图文", "视频"], width=8, state="readonly").pack(side=tk.LEFT, padx=(2, 15))
        
        ttk.Label(filter_row, text="时间范围:").pack(side=tk.LEFT)
        self.date_filter_var = tk.StringVar(value="全部")
        ttk.Combobox(filter_row, textvariable=self.date_filter_var,
                    values=["全部", "今天", "本周", "本月"], width=8, state="readonly").pack(side=tk.LEFT)
        
        # === 速度模式 ===
        speed_frame = ttk.LabelFrame(parent, text="速度模式", padding="10")
        speed_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.crawl_mode_var = tk.StringVar(value="standard")
        speed_row = ttk.Frame(speed_frame)
        speed_row.pack(fill=tk.X)
        
        ttk.Radiobutton(speed_row, text="标准模式（完整数据）", variable=self.crawl_mode_var, 
                       value="standard").pack(side=tk.LEFT, padx=(0, 15))
        ttk.Radiobutton(speed_row, text="快速模式（减少等待）", variable=self.crawl_mode_var, 
                       value="fast").pack(side=tk.LEFT, padx=(0, 15))
        ttk.Radiobutton(speed_row, text="极速模式（列表直取）", variable=self.crawl_mode_var, 
                       value="turbo").pack(side=tk.LEFT)
        
        # === 控制按钮 ===
        btn_frame = ttk.Frame(parent)
        btn_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.start_btn = ttk.Button(btn_frame, text="开始爬取", command=self._start_crawl, width=12)
        self.start_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        self.stop_btn = ttk.Button(btn_frame, text="停止", command=self._stop_crawl, state=tk.DISABLED, width=10)
        self.stop_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        ttk.Button(btn_frame, text="使用已保存Cookie", command=self._use_saved_cookies, width=18).pack(side=tk.LEFT, padx=(0, 5))
        
        ttk.Button(btn_frame, text="打开数据", command=self._open_data_dir, width=10).pack(side=tk.RIGHT)
        ttk.Button(btn_frame, text="打包图片", command=self._zip_images, width=10).pack(side=tk.RIGHT, padx=(0, 5))
        
        # === 进度区域 ===
        progress_frame = ttk.LabelFrame(parent, text="运行状态", padding="10")
        progress_frame.pack(fill=tk.X, pady=(0, 10))
        
        prog_row = ttk.Frame(progress_frame)
        prog_row.pack(fill=tk.X)
        self.total_progress = ttk.Progressbar(prog_row, length=400, mode='determinate')
        self.total_progress.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        self.progress_label = ttk.Label(prog_row, text="0%")
        self.progress_label.pack(side=tk.LEFT)
        
        stat_row = ttk.Frame(progress_frame)
        stat_row.pack(fill=tk.X, pady=5)
        
        self.status_var = tk.StringVar(value="就绪")
        ttk.Label(stat_row, text="状态:").pack(side=tk.LEFT)
        ttk.Label(stat_row, textvariable=self.status_var, foreground="blue").pack(side=tk.LEFT, padx=(5, 20))
        
        self.notes_var = tk.StringVar(value="笔记: 0")
        ttk.Label(stat_row, textvariable=self.notes_var).pack(side=tk.LEFT, padx=(0, 15))
        
        self.images_var = tk.StringVar(value="图片: 0")
        ttk.Label(stat_row, textvariable=self.images_var).pack(side=tk.LEFT, padx=(0, 15))
        
        self.videos_var = tk.StringVar(value="视频: 0")
        ttk.Label(stat_row, textvariable=self.videos_var).pack(side=tk.LEFT, padx=(0, 15))
        
        self.time_var = tk.StringVar(value="用时: 0秒")
        ttk.Label(stat_row, textvariable=self.time_var).pack(side=tk.LEFT)
        
        # === 日志区域 ===
        log_frame = ttk.LabelFrame(parent, text="运行日志", padding="5")
        log_frame.pack(fill=tk.BOTH, expand=True)
        
        self.log_text = scrolledtext.ScrolledText(log_frame, height=8, state=tk.DISABLED)
        self.log_text.pack(fill=tk.BOTH, expand=True)
        
        self.log_text.tag_config("INFO", foreground="black")
        self.log_text.tag_config("SUCCESS", foreground="green")
        self.log_text.tag_config("WARNING", foreground="orange")
        self.log_text.tag_config("ERROR", foreground="red")
    
    def _create_result_page(self, parent):
        """创建爬取结果展示页面"""
        # === 工具栏第一行 ===
        toolbar_frame = ttk.Frame(parent)
        toolbar_frame.pack(fill=tk.X, pady=(0, 5))
        
        # 左侧：数据源选择
        ttk.Label(toolbar_frame, text="数据源:").pack(side=tk.LEFT, padx=(0, 5))
        self.data_source_var = tk.StringVar(value="当前爬取")
        self.data_source_combo = ttk.Combobox(toolbar_frame, textvariable=self.data_source_var, 
                                               values=["当前爬取", "历史数据库"], width=12, state="readonly")
        self.data_source_combo.pack(side=tk.LEFT, padx=(0, 10))
        self.data_source_combo.bind("<<ComboboxSelected>>", self._on_data_source_change)
        
        # 关键词筛选
        ttk.Label(toolbar_frame, text="关键词:").pack(side=tk.LEFT, padx=(0, 5))
        self.filter_keyword_var = tk.StringVar()
        self.filter_keyword_entry = ttk.Entry(toolbar_frame, textvariable=self.filter_keyword_var, width=12)
        self.filter_keyword_entry.pack(side=tk.LEFT, padx=(0, 10))
        
        # 右侧按钮
        ttk.Button(toolbar_frame, text="刷新", command=self._refresh_results).pack(side=tk.RIGHT, padx=5)
        ttk.Button(toolbar_frame, text="导出Excel", command=self._export_results).pack(side=tk.RIGHT, padx=5)
        ttk.Button(toolbar_frame, text="删除选中", command=self._delete_selected).pack(side=tk.RIGHT, padx=5)
        ttk.Button(toolbar_frame, text="清空当前", command=self._clear_results).pack(side=tk.RIGHT, padx=5)
        
        # === 工具栏第二行：批次选择 + 搜索筛选 ===
        filter_frame = ttk.Frame(parent)
        filter_frame.pack(fill=tk.X, pady=(0, 5))
        
        # 左侧：批次选择
        ttk.Label(filter_frame, text="爬取批次:").pack(side=tk.LEFT, padx=(0, 5))
        self.crawl_batch_var = tk.StringVar(value="全部")
        self.crawl_batch_combo = ttk.Combobox(filter_frame, textvariable=self.crawl_batch_var, 
                                               width=30, state="readonly")
        self.crawl_batch_combo.pack(side=tk.LEFT, padx=(0, 5))
        self.crawl_batch_combo.bind("<<ComboboxSelected>>", self._on_batch_select)
        
        ttk.Button(filter_frame, text="刷新", command=self._refresh_crawl_batches, width=5).pack(side=tk.LEFT, padx=2)
        ttk.Button(filter_frame, text="删除", command=self._delete_batch_folder, width=5).pack(side=tk.LEFT, padx=2)
        ttk.Button(filter_frame, text="打开", command=self._open_batch_folder, width=5).pack(side=tk.LEFT, padx=2)
        
        # 分隔
        ttk.Separator(filter_frame, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=10)
        
        # 右侧：搜索筛选
        ttk.Label(filter_frame, text="搜索:").pack(side=tk.LEFT, padx=(0, 5))
        self.search_var = tk.StringVar()
        self.search_entry = ttk.Entry(filter_frame, textvariable=self.search_var, width=15)
        self.search_entry.pack(side=tk.LEFT, padx=(0, 5))
        self.search_entry.bind("<Return>", lambda e: self._filter_results())
        
        ttk.Button(filter_frame, text="筛选", command=self._filter_results, width=5).pack(side=tk.LEFT, padx=2)
        ttk.Button(filter_frame, text="重置", command=self._reset_filter, width=5).pack(side=tk.LEFT, padx=2)
        
        # 类型筛选
        ttk.Label(filter_frame, text="类型:").pack(side=tk.LEFT, padx=(10, 5))
        self.type_filter_var = tk.StringVar(value="全部")
        type_combo = ttk.Combobox(filter_frame, textvariable=self.type_filter_var, 
                                  values=["全部", "图文", "视频"], width=6, state="readonly")
        type_combo.pack(side=tk.LEFT)
        type_combo.bind("<<ComboboxSelected>>", lambda e: self._filter_results())
        
        # 初始化批次列表
        self._refresh_crawl_batches()
        
        # === 统计信息栏（带数据卡片）===
        stats_frame = ttk.Frame(parent)
        stats_frame.pack(fill=tk.X, pady=(0, 8))
        
        # 统计卡片样式
        self.result_count_label = ttk.Label(stats_frame, text="总计: 0 条", font=("", 9, "bold"))
        self.result_count_label.pack(side=tk.LEFT, padx=(0, 15))
        
        self.stats_image_label = ttk.Label(stats_frame, text="图文: 0", foreground="#2196F3")
        self.stats_image_label.pack(side=tk.LEFT, padx=(0, 15))
        
        self.stats_video_label = ttk.Label(stats_frame, text="视频: 0", foreground="#FF5722")
        self.stats_video_label.pack(side=tk.LEFT, padx=(0, 15))
        
        self.stats_likes_label = ttk.Label(stats_frame, text="总点赞: 0", foreground="#E91E63")
        self.stats_likes_label.pack(side=tk.LEFT, padx=(0, 15))
        
        # 导出按钮
        ttk.Button(stats_frame, text="导出Excel", command=self._quick_export_excel, width=10).pack(side=tk.RIGHT, padx=5)
        ttk.Button(stats_frame, text="复制全部", command=self._copy_all_data, width=8).pack(side=tk.RIGHT, padx=5)
        
        # === 主区域：左边表格，右边详情 ===
        main_paned = ttk.PanedWindow(parent, orient=tk.HORIZONTAL)
        main_paned.pack(fill=tk.BOTH, expand=True)
        
        # 左侧：结果表格
        left_frame = ttk.Frame(main_paned)
        main_paned.add(left_frame, weight=3)
        
        columns = ("序号", "类型", "标题", "作者", "点赞", "收藏", "评论")
        self.result_tree = ttk.Treeview(left_frame, columns=columns, show="headings", height=22)
        
        # 配置表格样式 - 斑马纹
        style = ttk.Style()
        style.configure("Treeview", rowheight=28, font=("", 9))
        style.configure("Treeview.Heading", font=("", 9, "bold"))
        self.result_tree.tag_configure('oddrow', background='#f8f8f8')
        self.result_tree.tag_configure('evenrow', background='#ffffff')
        self.result_tree.tag_configure('video', foreground='#FF5722')
        self.result_tree.tag_configure('image', foreground='#2196F3')
        
        # 表头 - 点击可排序
        for col in columns:
            self.result_tree.heading(col, text=col, command=lambda c=col: self._sort_by_column(c))
        
        self.result_tree.column("序号", width=40, anchor="center")
        self.result_tree.column("类型", width=50, anchor="center")
        self.result_tree.column("标题", width=220, anchor="w")
        self.result_tree.column("作者", width=90, anchor="w")
        self.result_tree.column("点赞", width=60, anchor="e")
        self.result_tree.column("收藏", width=60, anchor="e")
        self.result_tree.column("评论", width=60, anchor="e")
        
        # 滚动条
        scrollbar_y = ttk.Scrollbar(left_frame, orient=tk.VERTICAL, command=self.result_tree.yview)
        scrollbar_x = ttk.Scrollbar(left_frame, orient=tk.HORIZONTAL, command=self.result_tree.xview)
        self.result_tree.configure(yscrollcommand=scrollbar_y.set, xscrollcommand=scrollbar_x.set)
        
        self.result_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 右键菜单
        self.tree_context_menu = tk.Menu(self.result_tree, tearoff=0)
        self.tree_context_menu.add_command(label="复制标题", command=self._copy_title)
        self.tree_context_menu.add_command(label="复制作者", command=self._copy_author)
        self.tree_context_menu.add_command(label="复制链接", command=self._copy_link)
        self.tree_context_menu.add_separator()
        self.tree_context_menu.add_command(label="打开原文", command=self._open_note_link)
        self.tree_context_menu.add_command(label="打开文件夹", command=self._open_images_folder)
        self.tree_context_menu.add_separator()
        self.tree_context_menu.add_command(label="删除此条", command=self._delete_single_note)
        self.result_tree.bind("<Button-3>", self._show_tree_context_menu)
        
        # 右侧：详情面板
        right_frame = ttk.Frame(main_paned)
        main_paned.add(right_frame, weight=2)
        
        # 详情顶部：标题 + 快捷操作
        detail_header = ttk.Frame(right_frame)
        detail_header.pack(fill=tk.X, pady=(0, 5))
        
        self.detail_title_label = ttk.Label(detail_header, text="选择笔记查看详情", 
                                            font=("", 11, "bold"), wraplength=300)
        self.detail_title_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # 操作按钮（图标化）
        btn_frame = ttk.Frame(detail_header)
        btn_frame.pack(side=tk.RIGHT)
        ttk.Button(btn_frame, text="📂", command=self._open_images_folder, width=3).pack(side=tk.LEFT, padx=1)
        ttk.Button(btn_frame, text="▶", command=self._play_video, width=3).pack(side=tk.LEFT, padx=1)
        ttk.Button(btn_frame, text="🔗", command=self._open_note_link, width=3).pack(side=tk.LEFT, padx=1)
        ttk.Button(btn_frame, text="📋", command=self._copy_note_content, width=3).pack(side=tk.LEFT, padx=1)
        
        # 数据卡片区
        info_cards = ttk.Frame(right_frame)
        info_cards.pack(fill=tk.X, pady=(0, 8))
        
        # 互动数据展示
        self.detail_likes = ttk.Label(info_cards, text="❤ 0", foreground="#E91E63", font=("", 10))
        self.detail_likes.pack(side=tk.LEFT, padx=(0, 15))
        self.detail_collects = ttk.Label(info_cards, text="⭐ 0", foreground="#FF9800", font=("", 10))
        self.detail_collects.pack(side=tk.LEFT, padx=(0, 15))
        self.detail_comments = ttk.Label(info_cards, text="💬 0", foreground="#2196F3", font=("", 10))
        self.detail_comments.pack(side=tk.LEFT, padx=(0, 15))
        self.detail_author = ttk.Label(info_cards, text="", foreground="#666", font=("", 9))
        self.detail_author.pack(side=tk.RIGHT)
        
        # 详情内容（减小高度，给预览更多空间）
        self.detail_text = scrolledtext.ScrolledText(right_frame, height=8, state=tk.DISABLED, 
                                                     wrap=tk.WORD, font=("", 9))
        self.detail_text.pack(fill=tk.BOTH, expand=True, pady=(0, 5))
        
        # 图片预览区（更大）
        preview_frame = ttk.LabelFrame(right_frame, text="媒体预览", padding="5")
        preview_frame.pack(fill=tk.BOTH, expand=True)
        
        # 预览导航
        preview_nav = ttk.Frame(preview_frame)
        preview_nav.pack(fill=tk.X, pady=(0, 5))
        
        self.preview_page_label = ttk.Label(preview_nav, text="")
        self.preview_page_label.pack(side=tk.LEFT)
        
        ttk.Button(preview_nav, text="◀", command=self._prev_preview_page, width=3).pack(side=tk.RIGHT, padx=2)
        ttk.Button(preview_nav, text="▶", command=self._next_preview_page, width=3).pack(side=tk.RIGHT, padx=2)
        ttk.Button(preview_nav, text="查看大图", command=self._open_image_viewer, width=8).pack(side=tk.RIGHT, padx=5)
        
        self.preview_canvas = tk.Canvas(preview_frame, height=180, bg="#f0f0f0")
        self.preview_canvas.pack(fill=tk.BOTH, expand=True)
        self.preview_canvas.bind("<Double-Button-1>", self._on_preview_double_click)
        
        # 存储当前选中的笔记数据
        self.current_selected_note = None
        self.preview_image_paths = []
        self.preview_comment_images = []  # 评论图片路径
        self.preview_images = []  # 保持图片引用
        self.current_video_path = None  # 当前预览的视频路径
        self.preview_page = 0  # 预览分页
        self.preview_page_size = 5  # 每页显示数量
        self.sort_column = None  # 排序列
        self.sort_reverse = False  # 排序方向
        self.filtered_notes = []  # 筛选后的数据
        
        # 绑定事件
        self.result_tree.bind("<<TreeviewSelect>>", self._on_result_select)
        self.result_tree.bind("<Double-Button-1>", self._on_result_double_click)
    
    def _refresh_crawl_batches(self):
        """刷新爬取批次列表"""
        import glob
        batches = ["全部"]
        
        # 扫描images目录下的所有文件夹
        if os.path.exists("images"):
            folders = []
            for folder in os.listdir("images"):
                folder_path = os.path.join("images", folder)
                if os.path.isdir(folder_path):
                    # 获取文件夹信息
                    try:
                        mtime = os.path.getmtime(folder_path)
                        # 计算文件夹内的图片数量
                        img_count = len(glob.glob(f"{folder_path}/**/*.jpg", recursive=True))
                        img_count += len(glob.glob(f"{folder_path}/**/*.png", recursive=True))
                        folders.append((folder, mtime, img_count))
                    except:
                        folders.append((folder, 0, 0))
            
            # 按修改时间排序（最新的在前）
            folders.sort(key=lambda x: x[1], reverse=True)
            
            # 格式化显示
            from datetime import datetime
            for folder, mtime, count in folders:
                if mtime > 0:
                    time_str = datetime.fromtimestamp(mtime).strftime("%m-%d %H:%M")
                    batches.append(f"{folder} ({count}张) [{time_str}]")
                else:
                    batches.append(f"{folder} ({count}张)")
        
        self.crawl_batch_combo['values'] = batches
        if batches:
            self.crawl_batch_combo.current(0)
    
    def _on_batch_select(self, event=None):
        """选择爬取批次"""
        selected = self.crawl_batch_var.get()
        if selected == "全部":
            self._load_all_batch_images()
        else:
            # 提取文件夹名
            folder_name = selected.split(" (")[0]
            self._load_batch_images(folder_name)
    
    def _filter_results(self):
        """筛选结果"""
        search_text = self.search_var.get().strip().lower()
        type_filter = self.type_filter_var.get()
        
        # 获取数据源
        if self.data_source_var.get() == "历史数据库":
            source_notes = getattr(self, 'history_notes_data', [])
        else:
            source_notes = self.all_notes_data
        
        # 筛选
        filtered = []
        for note in source_notes:
            # 类型筛选
            note_type = note.get('note_type', '图文')
            if type_filter != "全部":
                if type_filter == "视频" and note_type != "视频":
                    continue
                if type_filter == "图文" and note_type == "视频":
                    continue
            
            # 文本搜索
            if search_text:
                title = (note.get('title', '') or '').lower()
                author = (note.get('author', '') or '').lower()
                content = (note.get('content', '') or '').lower()
                if search_text not in title and search_text not in author and search_text not in content:
                    continue
            
            filtered.append(note)
        
        self.filtered_notes = filtered
        self._refresh_table_with_notes(filtered)
    
    def _reset_filter(self):
        """重置筛选"""
        self.search_var.set("")
        self.type_filter_var.set("全部")
        self.filtered_notes = []
        self._on_data_source_change()
    
    def _refresh_table_with_notes(self, notes):
        """用指定数据刷新表格"""
        for item in self.result_tree.get_children():
            self.result_tree.delete(item)
        
        total_likes = 0
        image_count = 0
        video_count = 0
        
        for i, note in enumerate(notes):
            note_type = "视频" if note.get('note_type') == "视频" else "图文"
            like_count = note.get('like_count', 0) or 0
            collect_count = note.get('collect_count', 0) or 0
            comment_count = note.get('comment_count', 0) or 0
            
            # 斑马纹和类型颜色
            tags = ('oddrow',) if i % 2 else ('evenrow',)
            if note_type == "视频":
                tags = tags + ('video',)
                video_count += 1
            else:
                tags = tags + ('image',)
                image_count += 1
            
            try:
                total_likes += int(like_count) if str(like_count).isdigit() else 0
            except:
                pass
            
            self.result_tree.insert("", tk.END, values=(
                i + 1, note_type, 
                (note.get('title', '') or '')[:28],
                (note.get('author', '') or '')[:12],
                like_count, collect_count, comment_count
            ), tags=tags)
        
        # 更新统计
        self.result_count_label.config(text=f"总计: {len(notes)} 条")
        self.stats_image_label.config(text=f"图文: {image_count}")
        self.stats_video_label.config(text=f"视频: {video_count}")
        self.stats_likes_label.config(text=f"总点赞: {total_likes:,}")
    
    def _sort_by_column(self, col):
        """点击表头排序"""
        # 获取当前数据
        if self.filtered_notes:
            notes = self.filtered_notes
        elif self.data_source_var.get() == "历史数据库":
            notes = getattr(self, 'history_notes_data', [])
        else:
            notes = self.all_notes_data
        
        if not notes:
            return
        
        # 切换排序方向
        if self.sort_column == col:
            self.sort_reverse = not self.sort_reverse
        else:
            self.sort_column = col
            self.sort_reverse = False
        
        # 排序映射
        key_map = {
            "序号": lambda x: x.get('idx', 0) or 0,
            "类型": lambda x: x.get('note_type', ''),
            "标题": lambda x: x.get('title', '') or '',
            "作者": lambda x: x.get('author', '') or '',
            "点赞": lambda x: int(x.get('like_count', 0) or 0),
            "收藏": lambda x: int(x.get('collect_count', 0) or 0),
            "评论": lambda x: int(x.get('comment_count', 0) or 0),
        }
        
        key_func = key_map.get(col)
        if key_func:
            try:
                notes_sorted = sorted(notes, key=key_func, reverse=self.sort_reverse)
                self._refresh_table_with_notes(notes_sorted)
            except:
                pass
    
    def _show_tree_context_menu(self, event):
        """显示右键菜单"""
        item = self.result_tree.identify_row(event.y)
        if item:
            self.result_tree.selection_set(item)
            self.tree_context_menu.post(event.x_root, event.y_root)
    
    def _copy_title(self):
        """复制标题"""
        if self.current_selected_note:
            title = self.current_selected_note.get('title', '')
            self.root.clipboard_clear()
            self.root.clipboard_append(title)
    
    def _copy_author(self):
        """复制作者"""
        if self.current_selected_note:
            author = self.current_selected_note.get('author', '')
            self.root.clipboard_clear()
            self.root.clipboard_append(author)
    
    def _copy_link(self):
        """复制链接"""
        if self.current_selected_note:
            note_id = self.current_selected_note.get('note_id', '')
            if note_id:
                link = f"https://www.xiaohongshu.com/explore/{note_id}"
                self.root.clipboard_clear()
                self.root.clipboard_append(link)
    
    def _copy_note_content(self):
        """复制笔记内容"""
        if self.current_selected_note:
            note = self.current_selected_note
            content = f"{note.get('title', '')}\n\n作者: {note.get('author', '')}\n\n{note.get('content', '')}"
            self.root.clipboard_clear()
            self.root.clipboard_append(content)
    
    def _copy_all_data(self):
        """复制全部数据为文本"""
        items = self.result_tree.get_children()
        if not items:
            return
        
        lines = ["序号\t类型\t标题\t作者\t点赞\t收藏\t评论"]
        for item in items:
            values = self.result_tree.item(item)['values']
            lines.append("\t".join(str(v) for v in values))
        
        self.root.clipboard_clear()
        self.root.clipboard_append("\n".join(lines))
        messagebox.showinfo("成功", f"已复制 {len(items)} 条数据到剪贴板")
    
    def _quick_export_excel(self):
        """快速导出当前表格数据到Excel"""
        items = self.result_tree.get_children()
        if not items:
            messagebox.showinfo("提示", "没有数据可导出")
            return
        
        try:
            import pandas as pd
            data = []
            for item in items:
                values = self.result_tree.item(item)['values']
                data.append({
                    '序号': values[0], '类型': values[1], '标题': values[2],
                    '作者': values[3], '点赞': values[4], '收藏': values[5], '评论': values[6]
                })
            
            df = pd.DataFrame(data)
            filename = f"导出数据_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            filepath = os.path.join("images", filename)
            df.to_excel(filepath, index=False)
            messagebox.showinfo("成功", f"已导出到: {filepath}")
        except Exception as e:
            messagebox.showerror("错误", f"导出失败: {e}")
    
    def _delete_single_note(self):
        """删除单条笔记"""
        if not self.current_selected_note:
            return
        
        if not messagebox.askyesno("确认", "确定要删除这条笔记吗？"):
            return
        
        try:
            # 删除文件夹
            folder = self.current_selected_note.get('folder_path') or self.current_selected_note.get('path')
            if folder and os.path.exists(folder):
                import shutil
                shutil.rmtree(folder)
            
            # 刷新显示
            self._on_batch_select()
        except Exception as e:
            messagebox.showerror("错误", f"删除失败: {e}")
    
    def _prev_preview_page(self):
        """上一页预览"""
        if self.preview_page > 0:
            self.preview_page -= 1
            self._render_preview_page()
    
    def _next_preview_page(self):
        """下一页预览"""
        total = len(self.preview_image_paths)
        max_page = max(1, (total + self.preview_page_size - 1) // self.preview_page_size)
        if self.preview_page < max_page - 1:
            self.preview_page += 1
            self._render_preview_page()
    
    def _open_image_viewer(self):
        """打开图片查看器"""
        if not self.preview_image_paths:
            return
        
        try:
            # 打开第一张图片
            if self.preview_image_paths:
                os.startfile(self.preview_image_paths[0])
        except Exception as e:
            messagebox.showerror("错误", f"打开失败: {e}")
    
    def _load_batch_images(self, folder_name):
        """加载指定批次的图片"""
        import glob
        
        folder_path = os.path.abspath(os.path.join("images", folder_name))
        if not os.path.exists(folder_path):
            return
        
        # 清空表格
        for item in self.result_tree.get_children():
            self.result_tree.delete(item)
        
        # 扫描文件夹下的所有笔记
        note_folders = []
        for note_folder in os.listdir(folder_path):
            note_path = os.path.abspath(os.path.join(folder_path, note_folder))
            if os.path.isdir(note_path) and note_folder.startswith("note_"):
                # 使用绝对路径查找图片和视频
                images = [os.path.abspath(f) for f in glob.glob(os.path.join(note_path, "*.jpg"))]
                images += [os.path.abspath(f) for f in glob.glob(os.path.join(note_path, "*.png"))]
                images += [os.path.abspath(f) for f in glob.glob(os.path.join(note_path, "*.webp"))]
                videos = [os.path.abspath(f) for f in glob.glob(os.path.join(note_path, "*.mp4"))]
                if images or videos:
                    # 提取序号
                    try:
                        idx = int(note_folder.split("_")[1])
                    except:
                        idx = 0
                    note_folders.append({
                        'folder': note_folder,
                        'path': note_path,
                        'idx': idx,
                        'images': images,
                        'videos': videos,
                        'image_count': len(images),
                        'has_video': len(videos) > 0
                    })
        
        # 按序号排序
        note_folders.sort(key=lambda x: x['idx'])
        
        # 存储当前批次数据
        self.batch_notes_data = note_folders
        self.current_batch_folder = folder_path
        
        # 填充表格
        from datetime import datetime
        folder_time = ""
        try:
            mtime = os.path.getmtime(folder_path)
            folder_time = datetime.fromtimestamp(mtime).strftime("%m-%d %H:%M")
        except:
            pass
        
        for i, note in enumerate(note_folders):
            note_type = "视频" if note['has_video'] else "图文"
            self.result_tree.insert("", tk.END, values=(
                note['idx'],
                note_type,
                f"笔记{note['idx']}",
                f"{note['image_count']}张",
                "-",
                "-",
                "-"
            ))
        
        self.result_count_label.config(text=f"共 {len(note_folders)} 个笔记")
    
    def _load_all_batch_images(self):
        """加载所有批次的摘要"""
        import glob
        
        # 清空表格
        for item in self.result_tree.get_children():
            self.result_tree.delete(item)
        
        self.batch_notes_data = []
        self.current_batch_folder = None
        
        if not os.path.exists("images"):
            return
        
        folders = []
        for folder in os.listdir("images"):
            folder_path = os.path.join("images", folder)
            if os.path.isdir(folder_path):
                try:
                    mtime = os.path.getmtime(folder_path)
                    img_count = len(glob.glob(f"{folder_path}/**/*.jpg", recursive=True))
                    img_count += len(glob.glob(f"{folder_path}/**/*.png", recursive=True))
                    video_count = len(glob.glob(f"{folder_path}/**/*.mp4", recursive=True))
                    note_count = len([d for d in os.listdir(folder_path) if os.path.isdir(os.path.join(folder_path, d))])
                    folders.append({
                        'name': folder,
                        'path': folder_path,
                        'mtime': mtime,
                        'images': img_count,
                        'videos': video_count,
                        'notes': note_count
                    })
                except:
                    pass
        
        # 按时间排序
        folders.sort(key=lambda x: x['mtime'], reverse=True)
        
        from datetime import datetime
        for i, f in enumerate(folders):
            time_str = datetime.fromtimestamp(f['mtime']).strftime("%m-%d %H:%M")
            # 解析关键词
            keyword = f['name'].split("_")[0] if "_" in f['name'] else f['name']
            self.result_tree.insert("", tk.END, values=(
                i + 1,
                "批次",
                f"{keyword} ({time_str})",
                f"{f['notes']}笔记",
                f"{f['images']}图",
                f"{f['videos']}视频",
                "-"
            ))
        
        self.batch_notes_data = folders
        self.result_count_label.config(text=f"共 {len(folders)} 个爬取批次")
    
    def _delete_batch_folder(self):
        """删除选中的批次文件夹"""
        selected = self.crawl_batch_var.get()
        if selected == "全部":
            messagebox.showinfo("提示", "请先选择一个具体的爬取批次")
            return
        
        folder_name = selected.split(" (")[0]
        folder_path = os.path.join("images", folder_name)
        
        if not os.path.exists(folder_path):
            messagebox.showinfo("提示", "文件夹不存在")
            return
        
        # 计算内容
        import glob
        img_count = len(glob.glob(f"{folder_path}/**/*.jpg", recursive=True))
        img_count += len(glob.glob(f"{folder_path}/**/*.png", recursive=True))
        
        if not messagebox.askyesno("确认删除", 
            f"确定要删除整个爬取批次吗？\n\n文件夹: {folder_name}\n图片数量: {img_count}\n\n此操作不可恢复！"):
            return
        
        try:
            import shutil
            shutil.rmtree(folder_path)
            messagebox.showinfo("完成", f"已删除: {folder_name}")
            self._refresh_crawl_batches()
            self._load_all_batch_images()
        except Exception as e:
            messagebox.showerror("错误", f"删除失败: {e}")
    
    def _open_batch_folder(self):
        """打开批次文件夹"""
        selected = self.crawl_batch_var.get()
        if selected == "全部":
            if os.path.exists("images"):
                os.startfile(os.path.abspath("images"))
            return
        
        folder_name = selected.split(" (")[0]
        folder_path = os.path.join("images", folder_name)
        
        if os.path.exists(folder_path):
            os.startfile(os.path.abspath(folder_path))
        else:
            messagebox.showinfo("提示", "文件夹不存在")
    
    def _on_data_source_change(self, event=None):
        """切换数据源"""
        source = self.data_source_var.get()
        
        # 清空批次相关数据，避免影响其他视图
        self.batch_notes_data = []
        self.current_batch_folder = None
        
        if source == "历史数据库":
            self._load_history_data()
        else:
            self._show_current_data()
    
    def _get_date_filter(self):
        """获取日期过滤范围"""
        from datetime import datetime, timedelta
        
        date_filter = self.date_filter_var.get() if hasattr(self, 'date_filter_var') else "全部"
        
        if date_filter == "全部":
            return None, None
        
        now = datetime.now()
        
        if date_filter == "今天":
            start = now.replace(hour=0, minute=0, second=0, microsecond=0)
            end = start + timedelta(days=1)
        elif date_filter == "本周":
            start = now - timedelta(days=now.weekday())
            start = start.replace(hour=0, minute=0, second=0, microsecond=0)
            end = start + timedelta(days=7)
        elif date_filter == "本月":
            start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
            if now.month == 12:
                end = start.replace(year=now.year + 1, month=1)
            else:
                end = start.replace(month=now.month + 1)
        else:
            return None, None
        
        return start.strftime("%Y-%m-%d %H:%M:%S"), end.strftime("%Y-%m-%d %H:%M:%S")
    
    def _load_history_data(self):
        """从数据库加载历史数据"""
        try:
            keyword_filter = self.filter_keyword_var.get().strip()
            start_date, end_date = self._get_date_filter()
            
            conn = sqlite3.connect(self.config.db_path)
            cursor = conn.cursor()
            
            # 构建SQL查询
            conditions = []
            params = []
            
            if keyword_filter:
                conditions.append("keyword LIKE ?")
                params.append(f"%{keyword_filter}%")
            
            if start_date and end_date:
                conditions.append("crawl_time >= ? AND crawl_time < ?")
                params.extend([start_date, end_date])
            
            if conditions:
                sql = f"SELECT * FROM notes WHERE {' AND '.join(conditions)} ORDER BY crawl_time DESC"
                cursor.execute(sql, params)
            else:
                cursor.execute("SELECT * FROM notes ORDER BY crawl_time DESC LIMIT 500")
            
            rows = cursor.fetchall()
            columns = [desc[0] for desc in cursor.description]
            conn.close()
            
            # 清空表格
            for item in self.result_tree.get_children():
                self.result_tree.delete(item)
            
            # 临时存储历史数据
            self.history_notes_data = []
            for row in rows:
                note = dict(zip(columns, row))
                # 解析JSON字段
                try:
                    note['image_urls'] = json.loads(note.get('image_urls', '[]'))
                except:
                    note['image_urls'] = []
                try:
                    note['comments'] = json.loads(note.get('comments', '[]'))
                except:
                    note['comments'] = []
                self.history_notes_data.append(note)
            
            # 填充表格
            for i, note in enumerate(self.history_notes_data):
                note_type = "视频" if note.get('note_type') == "视频" else "图文"
                self.result_tree.insert("", tk.END, values=(
                    i + 1,
                    note_type,
                    note.get('title', '')[:25],
                    note.get('author', '')[:10],
                    note.get('like_count', 0),
                    note.get('collect_count', 0),
                    note.get('comment_count', 0)
                ))
            
            self.result_count_label.config(text=f"共 {len(self.history_notes_data)} 条历史记录")
            
        except Exception as e:
            messagebox.showerror("错误", f"加载历史数据失败: {e}")
    
    def _show_current_data(self):
        """显示当前爬取的数据"""
        from datetime import datetime
        for item in self.result_tree.get_children():
            self.result_tree.delete(item)
        
        for i, note in enumerate(self.all_notes_data):
            note_type = "视频" if note.get('note_type') == "视频" else "图文"
            self.result_tree.insert("", tk.END, values=(
                i + 1,
                note_type,
                note.get('title', '')[:25],
                note.get('author', '')[:8],
                note.get('like_count', 0),
                note.get('keyword', '')[:6],
                datetime.now().strftime("%m-%d %H:%M")
            ))
        
        self.result_count_label.config(text=f"共 {len(self.all_notes_data)} 条记录")
    
    def _refresh_results(self):
        """刷新结果"""
        self._on_data_source_change()
    
    def _add_result_to_table(self, note_data: dict, index: int):
        """添加一条结果到表格"""
        try:
            if self.data_source_var.get() != "当前爬取":
                return
            
            note_type = "视频" if note_data.get('note_type') == "视频" else "图文"
            like_count = note_data.get('like_count', 0)
            collect_count = note_data.get('collect_count', 0)
            comment_count = note_data.get('comment_count', 0)
            
            self.result_tree.insert("", tk.END, values=(
                index + 1,
                note_type,
                note_data.get('title', '')[:25],
                note_data.get('author', '')[:10],
                like_count,
                collect_count,
                comment_count
            ))
            
            count = len(self.result_tree.get_children())
            self.result_count_label.config(text=f"共 {count} 条记录")
            self.result_tree.see(self.result_tree.get_children()[-1])
        except Exception:
            pass
    
    def _on_result_select(self, event):
        """点击表格行显示详情"""
        try:
            selected = self.result_tree.selection()
            if not selected:
                return
            
            item = self.result_tree.item(selected[0])
            values = item['values']
            
            # 检查是否是批次视图（文件夹模式）
            data_source = self.data_source_var.get()
            batch_notes = getattr(self, 'batch_notes_data', [])
            batch_folder = getattr(self, 'current_batch_folder', None)
            
            # 检查用户是否明确选择了一个爬取批次（不是"全部"）
            selected_batch = getattr(self, 'crawl_batch_var', None)
            batch_selected = selected_batch and selected_batch.get() and selected_batch.get() != "全部"
            
            # 只有在没有明确选择批次的情况下，才根据数据源清空批次逻辑
            if not batch_selected:
                if data_source != "当前爬取" or (not batch_folder and self.all_notes_data):
                    batch_notes = []
                    batch_folder = None
            
            if batch_folder and batch_notes:
                # 批次内的笔记视图
                idx = int(values[0])
                for note in batch_notes:
                    if note.get('idx') == idx:
                        self.current_selected_note = note
                        
                        # 更新顶部信息
                        self.detail_title_label.config(text=f"笔记 {idx}")
                        self.detail_likes.config(text=f"❤ -")
                        self.detail_collects.config(text=f"⭐ -")
                        self.detail_comments.config(text=f"💬 -")
                        self.detail_author.config(text=f"{note.get('image_count', 0)}张图片")
                        
                        detail = f"文件夹: {note.get('folder', '')}\n"
                        detail += f"图片数量: {note.get('image_count', 0)}\n"
                        detail += f"视频: {'有' if note.get('has_video') else '无'}\n"
                        detail += f"\n--- 图片文件 ---\n"
                        for img in note.get('images', [])[:10]:
                            detail += f"• {os.path.basename(img)}\n"
                        
                        self.detail_text.config(state=tk.NORMAL)
                        self.detail_text.delete(1.0, tk.END)
                        self.detail_text.insert(tk.END, detail)
                        self.detail_text.config(state=tk.DISABLED)
                        
                        # 加载图片预览
                        self._load_batch_note_previews(note)
                        return
                return
            
            elif batch_notes and not batch_folder:
                # 全部批次视图
                index = int(values[0]) - 1
                if 0 <= index < len(batch_notes):
                    folder = batch_notes[index]
                    self.current_selected_note = {'folder_path': folder['path'], 'keyword': folder['name'].split("_")[0]}
                    
                    # 更新顶部信息
                    self.detail_title_label.config(text=folder['name'])
                    self.detail_likes.config(text=f"❤ -")
                    self.detail_collects.config(text=f"⭐ -")
                    self.detail_comments.config(text=f"💬 -")
                    self.detail_author.config(text=f"{folder.get('notes', 0)}个笔记")
                    
                    detail = f"笔记数量: {folder.get('notes', 0)}\n"
                    detail += f"图片数量: {folder.get('images', 0)}\n"
                    detail += f"视频数量: {folder.get('videos', 0)}\n"
                    
                    from datetime import datetime
                    mtime = folder.get('mtime', 0)
                    if mtime:
                        detail += f"创建时间: {datetime.fromtimestamp(mtime).strftime('%Y-%m-%d %H:%M:%S')}\n"
                    
                    detail += f"\n双击进入查看详细内容"
                    
                    self.detail_text.config(state=tk.NORMAL)
                    self.detail_text.delete(1.0, tk.END)
                    self.detail_text.insert(tk.END, detail)
                    self.detail_text.config(state=tk.DISABLED)
                    
                    # 清空预览
                    self.preview_canvas.delete("all")
                    self.preview_image_paths = []
                    self.preview_page_label.config(text="")
                    return
            
            # 原有逻辑：数据库或当前爬取
            index = int(values[0]) - 1
            
            # 使用筛选后的数据（如果有）
            if self.filtered_notes:
                notes = self.filtered_notes
            elif self.data_source_var.get() == "历史数据库":
                notes = getattr(self, 'history_notes_data', [])
            else:
                notes = self.all_notes_data
            
            if 0 <= index < len(notes):
                note = notes[index]
                self.current_selected_note = note
                
                # 更新顶部信息卡片
                title = note.get('title', '') or '无标题'
                self.detail_title_label.config(text=title[:40] + ('...' if len(title) > 40 else ''))
                self.detail_likes.config(text=f"❤ {note.get('like_count', 0)}")
                self.detail_collects.config(text=f"⭐ {note.get('collect_count', 0)}")
                self.detail_comments.config(text=f"💬 {note.get('comment_count', 0)}")
                self.detail_author.config(text=f"@{note.get('author', '')}")
                
                # 构建详情文本（简化版，因为主要信息已在上方显示）
                detail = f"类型: {note.get('note_type', '图文')}\n"
                detail += f"发布时间: {note.get('publish_time', '')}\n"
                if note.get('ip_region'):
                    detail += f"IP地区: {note.get('ip_region', '')}\n"
                detail += f"标签: {note.get('tags', '')}\n"
                detail += f"关键词: {note.get('keyword', '')}\n"
                detail += f"\n--- 内容 ---\n{note.get('content', '')}\n"
                
                # 评论 - 醒目显示
                comments = note.get('comments', [])
                if comments:
                    detail += f"\n{'='*40}\n"
                    detail += f">>> 热门评论 ({len(comments)}条) <<<\n"
                    detail += f"{'='*40}\n\n"
                    for i, c in enumerate(comments[:10], 1):
                        if isinstance(c, dict):
                            author = c.get('author', '') or c.get('user', '') or '匿名'
                            content = c.get('content', '')
                            time_str = c.get('time', '')
                            ip_str = c.get('ip', '')
                            likes = c.get('likes', 0)
                            has_image = c.get('has_image', False)
                            
                            # 构建评论头部信息
                            header = f"[{i}] @{author}"
                            if ip_str:
                                header += f" | {ip_str}"
                            if time_str:
                                header += f" | {time_str}"
                            if likes > 0:
                                header += f" | ❤️{likes}"
                            
                            detail += f"{header}\n"
                            
                            # 图片标记
                            if has_image:
                                detail += "    [含图片评论]\n"
                            
                            detail += f"    {content}\n\n"
                        else:
                            detail += f"[{i}] {c}\n\n"
                
                self.detail_text.config(state=tk.NORMAL)
                self.detail_text.delete(1.0, tk.END)
                self.detail_text.insert(tk.END, detail)
                self.detail_text.config(state=tk.DISABLED)
                
                # 加载图片预览
                self._load_image_previews(note)
                
        except Exception as e:
            print(f"选择错误: {e}")
    
    def _load_batch_note_previews(self, note):
        """加载批次笔记的图片预览（支持分页）"""
        self.preview_canvas.delete("all")
        self.preview_images = []
        
        # 获取所有图片路径（确保绝对路径）
        all_images = []
        for img_path in note.get('images', []):
            abs_path = os.path.abspath(img_path)
            if os.path.exists(abs_path):
                all_images.append(abs_path)
        
        # 检查视频
        videos = note.get('videos', [])
        self.current_video_path = None
        if videos:
            for v in videos:
                abs_v = os.path.abspath(v)
                if os.path.exists(abs_v):
                    self.current_video_path = abs_v
                    break
        
        # 保存所有图片路径（用于分页和大图查看）
        self.preview_image_paths = all_images
        self.preview_comment_images = []  # 批次模式无评论图片
        self.preview_page = 0  # 重置分页
        
        if not all_images and not self.current_video_path:
            self.preview_canvas.create_text(200, 75, text="暂无媒体文件", fill="#888")
            self.preview_page_label.config(text="")
            return
        
        # 使用通用的分页渲染
        self._render_preview_page()
    
    def _load_image_previews(self, note):
        """加载图片预览 - 只显示当前笔记的图片"""
        self.preview_canvas.delete("all")
        self.preview_images = []
        self.preview_image_paths = []
        
        import glob
        
        # 获取本地图片路径 - 优先使用存储的路径
        local_images = note.get('local_images', [])
        
        # 如果是字符串格式，转换为列表
        if isinstance(local_images, str):
            local_images = [p.strip() for p in local_images.split('|') if p.strip()]
        
        # 转换为绝对路径并验证
        valid_stored = []
        for p in local_images:
            if p:
                abs_p = os.path.abspath(p)
                if os.path.exists(abs_p):
                    valid_stored.append(abs_p)
        
        if valid_stored:
            local_images = valid_stored
        else:
            # 没有有效的存储路径，使用精确的批次目录+序号查找
            local_images = []
            batch_dir = note.get('batch_dir', '')
            
            # 获取当前选中的序号
            idx = None
            try:
                selected = self.result_tree.selection()
                if selected:
                    item = self.result_tree.item(selected[0])
                    idx = int(item['values'][0])
            except:
                pass
            
            # 只使用精确的批次目录+序号查找（不跨批次）
            if batch_dir and idx:
                abs_batch = os.path.abspath(batch_dir)
                # 精确匹配 note_{idx}_ 开头的文件夹
                pattern = f"{abs_batch}/note_{idx}_*/*.*"
                local_images = [os.path.abspath(f) for f in glob.glob(pattern) 
                               if f.lower().endswith(('.jpg', '.png', '.webp'))]
            
            # 不使用其他备用方法，避免跨笔记混淆
        
        # 过滤有效路径
        valid_images = [p for p in local_images if p and os.path.exists(p)]
        
        # 查找视频文件
        local_video = note.get('local_video', '')
        
        # 转换为绝对路径检查
        if local_video:
            abs_video = os.path.abspath(local_video)
            if not os.path.exists(abs_video):
                local_video = ""
            else:
                local_video = abs_video
        
        if not local_video:
            # 尝试从图片目录查找视频
            if valid_images:
                video_dir = os.path.dirname(os.path.abspath(valid_images[0]))
                video_path = os.path.join(video_dir, 'video.mp4')
                if os.path.exists(video_path):
                    local_video = video_path
            
            # 如果还没找到，使用 batch_dir 精确查找
            if not local_video:
                batch_dir = note.get('batch_dir', '')
                idx = None
                try:
                    selected = self.result_tree.selection()
                    if selected:
                        item = self.result_tree.item(selected[0])
                        idx = int(item['values'][0])
                except:
                    pass
                
                # 只使用 batch_dir 精确匹配，不跨批次查找
                if batch_dir and idx:
                    abs_batch = os.path.abspath(batch_dir)
                    pattern = f"{abs_batch}/note_{idx}_*/video.mp4"
                    videos = glob.glob(pattern)
                    if videos:
                        local_video = os.path.abspath(videos[0])
        
        self.current_video_path = local_video if local_video and os.path.exists(local_video) else None
        
        # 查找评论图片（在 comments 子文件夹中）
        comment_images = []
        if valid_images:
            note_dir = os.path.dirname(valid_images[0])
            comments_dir = os.path.join(note_dir, 'comments')
            if os.path.exists(comments_dir):
                comment_images = [os.path.abspath(f) for f in glob.glob(f"{comments_dir}/*.*")
                                  if f.lower().endswith(('.jpg', '.png', '.webp'))]
        elif batch_dir and idx:
            abs_batch = os.path.abspath(batch_dir)
            pattern = f"{abs_batch}/note_{idx}_*/comments/*.*"
            comment_images = [os.path.abspath(f) for f in glob.glob(pattern)
                              if f.lower().endswith(('.jpg', '.png', '.webp'))]
        
        self.preview_comment_images = comment_images  # 保存评论图片路径
        
        if not valid_images and not self.current_video_path and not comment_images:
            self.preview_canvas.create_text(200, 90, text="暂无本地媒体", fill="#888")
            self.preview_page_label.config(text="")
            return
        
        # 保存所有图片路径（用于分页）
        self.preview_image_paths = valid_images
        self.preview_page = 0  # 重置分页
        
        # 使用分页显示
        self._render_preview_page()
    
    def _render_preview_page(self):
        """渲染当前预览页"""
        self.preview_canvas.delete("all")
        self.preview_images = []
        
        total = len(self.preview_image_paths)
        has_video = self.current_video_path is not None
        
        # 计算分页
        items_per_page = self.preview_page_size
        if has_video and self.preview_page == 0:
            items_per_page -= 1  # 第一页留一个位置给视频
        
        start = self.preview_page * self.preview_page_size
        if has_video and self.preview_page == 0:
            start = 0
        elif has_video:
            start = (self.preview_page * self.preview_page_size) - 1
        
        end = min(start + items_per_page, total)
        page_images = self.preview_image_paths[start:end] if total > 0 else []
        
        # 获取评论图片数量
        comment_count = len(getattr(self, 'preview_comment_images', []))
        
        # 更新分页标签
        if total > 0 or has_video or comment_count > 0:
            max_page = max(1, (total + self.preview_page_size - 1) // self.preview_page_size)
            media_info = f"{total}张图片"
            if has_video:
                media_info += " + 视频"
            if comment_count > 0:
                media_info += f" + {comment_count}张评论图"
            self.preview_page_label.config(text=f"第{self.preview_page + 1}/{max_page}页 ({media_info})")
        else:
            self.preview_page_label.config(text="")
        
        try:
            from PIL import Image, ImageTk, ImageDraw
            x_offset = 10
            thumb_size = 145  # 增大缩略图尺寸
            
            # 第一页先显示视频缩略图
            if has_video and self.preview_page == 0:
                try:
                    video_thumb = Image.new('RGB', (thumb_size, thumb_size), color=(35, 35, 35))
                    draw = ImageDraw.Draw(video_thumb)
                    
                    # 绘制播放按钮
                    cx, cy = thumb_size // 2, thumb_size // 2 - 10
                    r = 30
                    play_points = [(cx - r//2, cy - r), (cx - r//2, cy + r), (cx + r, cy)]
                    draw.polygon(play_points, fill=(255, 255, 255))
                    draw.text((thumb_size//2 - 25, thumb_size - 25), "VIDEO", fill=(180, 180, 180))
                    
                    photo = ImageTk.PhotoImage(video_thumb)
                    self.preview_images.append(photo)
                    self.preview_canvas.create_image(x_offset, 10, anchor="nw", image=photo, tags="video_thumb")
                    x_offset += thumb_size + 10
                except:
                    pass
            
            # 显示笔记图片
            for i, img_path in enumerate(page_images):
                try:
                    img = Image.open(img_path)
                    img.thumbnail((thumb_size, thumb_size))
                    photo = ImageTk.PhotoImage(img)
                    self.preview_images.append(photo)
                    self.preview_canvas.create_image(x_offset, 10, anchor="nw", image=photo, tags=f"img_{i}")
                    x_offset += thumb_size + 10
                except:
                    continue
            
            # 如果是最后一页且有评论图片，显示评论图片（带蓝色边框标记）
            comment_images = getattr(self, 'preview_comment_images', [])
            total_pages = max(1, (total + self.preview_page_size - 1) // self.preview_page_size)
            is_last_page = (self.preview_page >= total_pages - 1)
            
            if is_last_page and comment_images and x_offset < 700:
                # 添加分隔线
                self.preview_canvas.create_line(x_offset + 5, 10, x_offset + 5, thumb_size + 10, fill="#2196F3", width=2)
                x_offset += 15
                
                # 显示评论图片（最多显示剩余空间能容纳的数量）
                remaining_slots = (700 - x_offset) // (thumb_size + 10)
                for i, img_path in enumerate(comment_images[:remaining_slots]):
                    try:
                        img = Image.open(img_path)
                        img.thumbnail((thumb_size - 6, thumb_size - 6))
                        
                        # 创建带蓝色边框的图片
                        bordered = Image.new('RGB', (thumb_size, thumb_size), color=(33, 150, 243))  # 蓝色边框
                        bordered.paste(img, (3, 3))
                        
                        photo = ImageTk.PhotoImage(bordered)
                        self.preview_images.append(photo)
                        self.preview_canvas.create_image(x_offset, 10, anchor="nw", image=photo, tags=f"comment_img_{i}")
                        x_offset += thumb_size + 10
                    except:
                        continue
            
            # 绑定点击事件
            self.preview_canvas.bind("<Button-1>", self._on_preview_click_with_video)
            
            if not self.preview_images:
                self.preview_canvas.create_text(200, 90, text="媒体加载失败", fill="#888")
        except ImportError:
            self.preview_canvas.create_text(200, 90, text="需要安装Pillow: pip install Pillow", fill="#888")
        except Exception as e:
            self.preview_canvas.create_text(200, 90, text=f"加载预览失败: {e}", fill="#888")
    
    def _on_preview_click(self, event):
        """点击预览图打开查看器"""
        if not self.preview_image_paths:
            return
        
        # 计算点击的是哪张图片
        x = event.x
        img_index = x // 140
        if 0 <= img_index < len(self.preview_image_paths):
            self._open_image_viewer(img_index)
    
    def _on_preview_click_with_video(self, event):
        """点击预览图或视频缩略图"""
        x = event.x
        slot_index = x // 140
        
        # 如果有视频，第一个位置是视频
        if self.current_video_path:
            if slot_index == 0:
                # 点击了视频缩略图，播放视频
                try:
                    abs_path = os.path.abspath(self.current_video_path)
                    if os.path.exists(abs_path):
                        os.startfile(abs_path)
                    else:
                        messagebox.showerror("错误", f"视频文件不存在: {abs_path}")
                except Exception as e:
                    messagebox.showerror("错误", f"无法播放视频: {e}")
                return
            else:
                # 点击了图片，调整索引（减1因为第一个是视频）
                img_index = slot_index - 1
                if 0 <= img_index < len(self.preview_image_paths):
                    self._open_image_viewer(img_index)
        else:
            # 没有视频，直接按图片索引处理
            if 0 <= slot_index < len(self.preview_image_paths):
                self._open_image_viewer(slot_index)
    
    def _open_image_viewer(self, start_index=0):
        """打开图片查看器"""
        if not self.preview_image_paths:
            return
        
        from PIL import Image, ImageTk
        
        # 创建查看器窗口
        viewer = tk.Toplevel(self.root)
        viewer.title("图片查看器")
        viewer.geometry("900x700")
        viewer.configure(bg="#1a1a1a")
        
        # 当前图片索引
        current_index = [start_index]
        photo_ref = [None]  # 保持图片引用
        
        # 顶部信息栏
        info_frame = tk.Frame(viewer, bg="#1a1a1a")
        info_frame.pack(fill=tk.X, pady=5)
        
        info_label = tk.Label(info_frame, text="", fg="white", bg="#1a1a1a", font=("", 10))
        info_label.pack()
        
        # 图片显示区域
        canvas = tk.Canvas(viewer, bg="#1a1a1a", highlightthickness=0)
        canvas.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # 底部按钮栏
        btn_frame = tk.Frame(viewer, bg="#1a1a1a")
        btn_frame.pack(fill=tk.X, pady=10)
        
        def update_image():
            idx = current_index[0]
            if 0 <= idx < len(self.preview_image_paths):
                img_path = self.preview_image_paths[idx]
                try:
                    img = Image.open(img_path)
                    
                    # 计算缩放尺寸（保持比例，适应窗口）
                    canvas_w = canvas.winfo_width() or 880
                    canvas_h = canvas.winfo_height() or 550
                    
                    img_w, img_h = img.size
                    ratio = min(canvas_w / img_w, canvas_h / img_h, 1.0)  # 不放大超过原尺寸
                    new_w = int(img_w * ratio)
                    new_h = int(img_h * ratio)
                    
                    if ratio < 1.0:
                        img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
                    
                    photo_ref[0] = ImageTk.PhotoImage(img)
                    
                    canvas.delete("all")
                    canvas.create_image(canvas_w // 2, canvas_h // 2, anchor="center", image=photo_ref[0])
                    
                    # 更新信息
                    filename = os.path.basename(img_path)
                    info_label.config(text=f"{idx + 1} / {len(self.preview_image_paths)}  |  {filename}  |  {img_w}x{img_h}")
                    
                except Exception as e:
                    canvas.delete("all")
                    canvas.create_text(440, 275, text=f"加载失败: {e}", fill="white")
        
        def prev_image():
            if current_index[0] > 0:
                current_index[0] -= 1
                update_image()
        
        def next_image():
            if current_index[0] < len(self.preview_image_paths) - 1:
                current_index[0] += 1
                update_image()
        
        def open_in_explorer():
            if 0 <= current_index[0] < len(self.preview_image_paths):
                path = self.preview_image_paths[current_index[0]]
                folder = os.path.dirname(os.path.abspath(path))
                os.startfile(folder)
        
        # 按钮
        ttk.Button(btn_frame, text="< 上一张", command=prev_image, width=12).pack(side=tk.LEFT, padx=20)
        ttk.Button(btn_frame, text="打开文件夹", command=open_in_explorer, width=12).pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="下一张 >", command=next_image, width=12).pack(side=tk.LEFT, padx=20)
        ttk.Button(btn_frame, text="关闭", command=viewer.destroy, width=10).pack(side=tk.RIGHT, padx=20)
        
        # 键盘绑定
        viewer.bind("<Left>", lambda e: prev_image())
        viewer.bind("<Right>", lambda e: next_image())
        viewer.bind("<Escape>", lambda e: viewer.destroy())
        viewer.bind("<space>", lambda e: next_image())
        
        # 窗口大小变化时重新加载图片
        def on_resize(event):
            if event.widget == canvas:
                viewer.after(100, update_image)
        canvas.bind("<Configure>", on_resize)
        
        # 初始显示
        viewer.after(50, update_image)
        
        # 居中显示
        viewer.update_idletasks()
        x = (viewer.winfo_screenwidth() - 900) // 2
        y = (viewer.winfo_screenheight() - 700) // 2
        viewer.geometry(f"900x700+{x}+{y}")
        
        viewer.focus_set()
    
    def _on_preview_double_click(self, event):
        """双击预览图打开查看器"""
        self._open_image_viewer(0)
    
    def _on_result_double_click(self, event):
        """双击表格行"""
        # 如果在全部批次视图，双击进入该批次
        batch_notes = getattr(self, 'batch_notes_data', [])
        batch_folder = getattr(self, 'current_batch_folder', None)
        
        if batch_notes and not batch_folder:
            # 全部批次视图，双击进入
            selected = self.result_tree.selection()
            if selected:
                item = self.result_tree.item(selected[0])
                index = int(item['values'][0]) - 1
                if 0 <= index < len(batch_notes):
                    folder = batch_notes[index]
                    folder_name = folder['name']
                    # 更新下拉框选择
                    for val in self.crawl_batch_combo['values']:
                        if val.startswith(folder_name):
                            self.crawl_batch_var.set(val)
                            self._load_batch_images(folder_name)
                            return
        else:
            # 其他视图，打开图片文件夹
            self._open_images_folder()
    
    def _open_images_folder(self):
        """打开图片文件夹"""
        if not self.current_selected_note:
            return
        
        local_images = self.current_selected_note.get('local_images', [])
        if local_images and os.path.exists(local_images[0]):
            folder = os.path.dirname(os.path.abspath(local_images[0]))
            os.startfile(folder)
        else:
            # 尝试根据关键词找目录
            keyword = self.current_selected_note.get('keyword', '')
            folder = f"images/{keyword}"
            if os.path.exists(folder):
                os.startfile(os.path.abspath(folder))
            else:
                messagebox.showinfo("提示", "未找到本地图片文件夹")
    
    def _play_video(self):
        """播放视频"""
        if not self.current_selected_note:
            return
        
        local_video = self.current_selected_note.get('local_video', '')
        if local_video:
            # 转换为绝对路径
            abs_path = os.path.abspath(local_video)
            if os.path.exists(abs_path):
                try:
                    os.startfile(abs_path)
                except Exception as e:
                    messagebox.showerror("错误", f"无法播放视频: {e}")
            else:
                messagebox.showinfo("提示", f"视频文件不存在: {abs_path}")
        else:
            video_url = self.current_selected_note.get('video_url', '')
            if video_url:
                import webbrowser
                webbrowser.open(video_url)
            else:
                messagebox.showinfo("提示", "该笔记没有视频")
    
    def _open_note_link(self):
        """打开笔记链接"""
        if not self.current_selected_note:
            return
        
        link = self.current_selected_note.get('note_link', '')
        if link:
            import webbrowser
            webbrowser.open(link)
        else:
            messagebox.showinfo("提示", "没有笔记链接")
    
    def _delete_crawl_batch(self):
        """删除整个爬取批次（按关键词和时间删除）"""
        if not self.current_selected_note:
            messagebox.showinfo("提示", "请先选择一条记录")
            return
        
        keyword = self.current_selected_note.get('keyword', '')
        crawl_time = self.current_selected_note.get('crawl_time', '')
        
        if not keyword:
            messagebox.showinfo("提示", "无法确定爬取批次")
            return
        
        # 查找该批次的所有记录
        try:
            conn = sqlite3.connect(self.config.db_path)
            cursor = conn.cursor()
            
            # 按关键词和爬取日期查找
            if crawl_time:
                date_prefix = crawl_time[:10]  # YYYY-MM-DD
                cursor.execute("SELECT COUNT(*) FROM notes WHERE keyword = ? AND crawl_time LIKE ?", 
                              (keyword, f"{date_prefix}%"))
            else:
                cursor.execute("SELECT COUNT(*) FROM notes WHERE keyword = ?", (keyword,))
            
            count = cursor.fetchone()[0]
            conn.close()
            
            if count == 0:
                messagebox.showinfo("提示", "未找到该批次的记录")
                return
            
            # 确认删除
            batch_info = f"关键词: {keyword}"
            if crawl_time:
                batch_info += f"\n日期: {crawl_time[:10]}"
            
            if not messagebox.askyesno("确认删除批次", 
                f"确定要删除整个爬取批次吗？\n\n{batch_info}\n共 {count} 条记录\n\n这将同时删除：\n• 数据库中的所有相关记录\n• 对应的图片/视频文件夹"):
                return
            
            # 删除数据库记录
            conn = sqlite3.connect(self.config.db_path)
            cursor = conn.cursor()
            if crawl_time:
                date_prefix = crawl_time[:10]
                cursor.execute("DELETE FROM notes WHERE keyword = ? AND crawl_time LIKE ?", 
                              (keyword, f"{date_prefix}%"))
            else:
                cursor.execute("DELETE FROM notes WHERE keyword = ?", (keyword,))
            deleted_db = cursor.rowcount
            conn.commit()
            conn.close()
            
            # 删除文件夹
            deleted_folders = 0
            import shutil
            import glob
            
            # 查找匹配的文件夹
            if crawl_time:
                # 新格式: images/{keyword}_{YYYYMMDD}_*
                date_str = crawl_time[:10].replace("-", "")
                pattern = f"images/{keyword}_{date_str}*"
            else:
                # 旧格式或所有该关键词的文件夹
                pattern = f"images/{keyword}*"
            
            for folder in glob.glob(pattern):
                if os.path.isdir(folder):
                    try:
                        shutil.rmtree(folder)
                        deleted_folders += 1
                    except Exception as e:
                        print(f"删除文件夹失败: {folder}, {e}")
            
            # 刷新显示
            self._load_history_data()
            
            messagebox.showinfo("完成", f"已删除批次:\n• 数据库记录: {deleted_db} 条\n• 文件夹: {deleted_folders} 个")
            
        except Exception as e:
            messagebox.showerror("错误", f"删除失败: {e}")
    
    def _clear_results(self):
        """清空当前结果"""
        if self.data_source_var.get() == "当前爬取":
            self.all_notes_data = []
        
        for item in self.result_tree.get_children():
            self.result_tree.delete(item)
        self.result_count_label.config(text="共 0 条记录")
        self.detail_text.config(state=tk.NORMAL)
        self.detail_text.delete(1.0, tk.END)
        self.detail_text.config(state=tk.DISABLED)
        self.preview_canvas.delete("all")
        self.current_selected_note = None
    
    def _delete_selected(self):
        """删除选中的记录"""
        selected = self.result_tree.selection()
        if not selected:
            messagebox.showinfo("提示", "请先选择要删除的记录")
            return
        
        # 确认删除
        count = len(selected)
        if not messagebox.askyesno("确认删除", f"确定要删除选中的 {count} 条记录吗？\n\n这将同时删除：\n• 数据库中的记录\n• 对应的本地图片/视频文件"):
            return
        
        deleted_count = 0
        
        for item_id in selected:
            try:
                item = self.result_tree.item(item_id)
                index = int(item['values'][0]) - 1
                
                # 获取笔记数据
                if self.data_source_var.get() == "历史数据库":
                    notes = getattr(self, 'history_notes_data', [])
                else:
                    notes = self.all_notes_data
                
                if 0 <= index < len(notes):
                    note = notes[index]
                    
                    # 1. 从数据库删除
                    note_id = note.get('note_id', '')
                    note_link = note.get('note_link', '')
                    if note_id or note_link:
                        try:
                            conn = sqlite3.connect(self.config.db_path)
                            cursor = conn.cursor()
                            if note_id:
                                cursor.execute("DELETE FROM notes WHERE note_id = ?", (note_id,))
                            elif note_link:
                                cursor.execute("DELETE FROM notes WHERE note_link = ?", (note_link,))
                            conn.commit()
                            conn.close()
                        except Exception:
                            pass
                    
                    # 2. 删除本地文件
                    local_images = note.get('local_images', [])
                    if local_images:
                        if isinstance(local_images, str):
                            local_images = local_images.split(' | ')
                        for img_path in local_images:
                            if img_path and os.path.exists(img_path):
                                try:
                                    # 删除图片所在文件夹
                                    folder = os.path.dirname(img_path)
                                    if folder and os.path.exists(folder):
                                        import shutil
                                        shutil.rmtree(folder)
                                        break  # 文件夹已删除，不需要继续
                                except Exception:
                                    pass
                    
                    # 删除视频文件
                    local_video = note.get('local_video', '')
                    if local_video and os.path.exists(local_video):
                        try:
                            os.remove(local_video)
                        except Exception:
                            pass
                    
                    deleted_count += 1
                
                # 从表格删除
                self.result_tree.delete(item_id)
                
            except Exception as e:
                print(f"删除失败: {e}")
                continue
        
        # 更新数据列表
        if self.data_source_var.get() == "历史数据库":
            self._load_history_data()  # 重新加载
        else:
            # 从当前数据中移除已删除的项
            remaining_indices = set()
            for item_id in self.result_tree.get_children():
                item = self.result_tree.item(item_id)
                remaining_indices.add(int(item['values'][0]) - 1)
            self.all_notes_data = [n for i, n in enumerate(self.all_notes_data) if i in remaining_indices]
        
        # 更新统计
        count = len(self.result_tree.get_children())
        self.result_count_label.config(text=f"共 {count} 条记录")
        
        messagebox.showinfo("完成", f"已删除 {deleted_count} 条记录")
    
    def _export_results(self):
        """导出结果到Excel"""
        if self.data_source_var.get() == "历史数据库":
            data = getattr(self, 'history_notes_data', [])
        else:
            data = self.all_notes_data
        
        if not data:
            messagebox.showwarning("提示", "没有数据可导出")
            return
        try:
            filepath = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel文件", "*.xlsx")],
                initialfile=f"爬取结果_{int(time.time())}.xlsx"
            )
            if filepath:
                df = pd.DataFrame(data)
                df.to_excel(filepath, index=False)
                messagebox.showinfo("成功", f"已导出到: {filepath}")
        except Exception as e:
            messagebox.showerror("错误", f"导出失败: {e}")
    
    def _create_content_page(self, parent):
        """创建内容选项页面"""
        # === 基础内容 ===
        basic_frame = ttk.LabelFrame(parent, text="基础内容", padding="10")
        basic_frame.pack(fill=tk.X, pady=(0, 10))
        
        row1 = ttk.Frame(basic_frame)
        row1.pack(fill=tk.X, pady=2)
        
        self.get_content_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(row1, text="获取笔记正文内容", variable=self.get_content_var).pack(side=tk.LEFT, padx=(0, 20))
        
        self.get_tags_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(row1, text="提取话题标签 (#xxx)", variable=self.get_tags_var).pack(side=tk.LEFT, padx=(0, 20))
        
        self.get_time_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(row1, text="获取发布时间", variable=self.get_time_var).pack(side=tk.LEFT)
        
        row2 = ttk.Frame(basic_frame)
        row2.pack(fill=tk.X, pady=2)
        
        self.get_interactions_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(row2, text="获取互动数据（点赞/收藏/评论数）", variable=self.get_interactions_var).pack(side=tk.LEFT)
        
        # === 图片视频 ===
        media_frame = ttk.LabelFrame(parent, text="图片/视频", padding="10")
        media_frame.pack(fill=tk.X, pady=(0, 10))
        
        row3 = ttk.Frame(media_frame)
        row3.pack(fill=tk.X, pady=2)
        
        self.download_images_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(row3, text="下载图片", variable=self.download_images_var).pack(side=tk.LEFT, padx=(0, 20))
        
        self.get_all_images_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(row3, text="获取全部图片（切换轮播）", variable=self.get_all_images_var).pack(side=tk.LEFT, padx=(0, 20))
        
        self.download_videos_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(row3, text="下载视频", variable=self.download_videos_var).pack(side=tk.LEFT)
        
        # === 评论 ===
        comment_frame = ttk.LabelFrame(parent, text="评论爬取", padding="10")
        comment_frame.pack(fill=tk.X, pady=(0, 10))
        
        row4 = ttk.Frame(comment_frame)
        row4.pack(fill=tk.X, pady=2)
        
        self.get_comments_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(row4, text="获取热门评论", variable=self.get_comments_var).pack(side=tk.LEFT, padx=(0, 20))
        
        ttk.Label(row4, text="评论数量:").pack(side=tk.LEFT)
        self.comments_count_var = tk.StringVar(value="10")
        ttk.Spinbox(row4, from_=1, to=50, textvariable=self.comments_count_var, width=6).pack(side=tk.LEFT, padx=5)
        
        # === 导出格式 ===
        export_frame = ttk.LabelFrame(parent, text="导出设置", padding="10")
        export_frame.pack(fill=tk.X, pady=(0, 10))
        
        row5 = ttk.Frame(export_frame)
        row5.pack(fill=tk.X, pady=2)
        
        ttk.Label(row5, text="导出格式:").pack(side=tk.LEFT)
        self.export_format_var = tk.StringVar(value="xlsx")
        ttk.Combobox(row5, textvariable=self.export_format_var,
                    values=["xlsx", "csv", "json"], width=10, state="readonly").pack(side=tk.LEFT, padx=(5, 20))
        
        self.export_db_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(row5, text="同时保存到SQLite数据库", variable=self.export_db_var).pack(side=tk.LEFT)
        
        # === 快捷预设 ===
        preset_frame = ttk.LabelFrame(parent, text="快捷预设", padding="10")
        preset_frame.pack(fill=tk.X, pady=(0, 10))
        
        preset_row = ttk.Frame(preset_frame)
        preset_row.pack(fill=tk.X)
        
        ttk.Button(preset_row, text="极速采集", command=self._preset_turbo, width=12).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(preset_row, text="完整数据", command=self._preset_complete, width=12).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(preset_row, text="只下图片", command=self._preset_images, width=12).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(preset_row, text="只下视频", command=self._preset_videos, width=12).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(preset_row, text="只要文本", command=self._preset_text, width=12).pack(side=tk.LEFT)
    
    def _create_analysis_page(self, parent):
        """创建数据分析页面"""
        # === 分析工具 ===
        tools_frame = ttk.LabelFrame(parent, text="分析工具", padding="10")
        tools_frame.pack(fill=tk.X, pady=(0, 10))
        
        row1 = ttk.Frame(tools_frame)
        row1.pack(fill=tk.X, pady=5)
        
        ttk.Button(row1, text="生成统计图表", command=self._generate_charts, width=16).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(row1, text="生成词云", command=self._generate_wordcloud, width=16).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(row1, text="生成分析报告", command=self._generate_report, width=16).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(row1, text="合并所有数据", command=self._merge_data, width=16).pack(side=tk.LEFT)
        
        # === 统计仪表盘 ===
        dashboard_frame = ttk.LabelFrame(parent, text="统计仪表盘", padding="10")
        dashboard_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # 统计卡片网格
        stats_grid = ttk.Frame(dashboard_frame)
        stats_grid.pack(fill=tk.X, pady=10)
        
        self.dashboard_labels = {}
        stats_items = [
            ("total_notes", "总笔记", "0"),
            ("total_likes", "总点赞", "0"),
            ("avg_likes", "平均点赞", "0"),
            ("max_likes", "最高点赞", "0"),
            ("total_collects", "总收藏", "0"),
            ("total_comments", "总评论", "0"),
            ("image_notes", "图文笔记", "0"),
            ("video_notes", "视频笔记", "0"),
        ]
        
        for i, (key, label, default) in enumerate(stats_items):
            row = i // 4
            col = i % 4
            
            card = ttk.Frame(stats_grid, relief="solid", borderwidth=1)
            card.grid(row=row, column=col, padx=10, pady=5, sticky="nsew")
            
            ttk.Label(card, text=label, font=("", 9)).pack(pady=(5, 0))
            self.dashboard_labels[key] = ttk.Label(card, text=default, font=("", 14, "bold"), foreground="#667eea")
            self.dashboard_labels[key].pack(pady=(0, 5))
        
        for i in range(4):
            stats_grid.columnconfigure(i, weight=1)
        
        # === 历史记录 ===
        history_frame = ttk.LabelFrame(parent, text="历史记录", padding="10")
        history_frame.pack(fill=tk.BOTH, expand=True)
        
        columns = ("时间", "关键词", "笔记数", "图片数", "文件")
        self.history_tree = ttk.Treeview(history_frame, columns=columns, show="headings", height=8)
        
        for col in columns:
            self.history_tree.heading(col, text=col)
            self.history_tree.column(col, width=100)
        
        scrollbar = ttk.Scrollbar(history_frame, orient=tk.VERTICAL, command=self.history_tree.yview)
        self.history_tree.configure(yscrollcommand=scrollbar.set)
        
        self.history_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 刷新历史
        self._refresh_history()
    
    def _create_settings_page(self, parent):
        """创建设置页面"""
        # === Cookie管理 ===
        cookie_frame = ttk.LabelFrame(parent, text="Cookie管理", padding="10")
        cookie_frame.pack(fill=tk.X, pady=(0, 10))
        
        row1 = ttk.Frame(cookie_frame)
        row1.pack(fill=tk.X, pady=2)
        
        self.save_cookies_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(row1, text="登录后自动保存Cookie", variable=self.save_cookies_var).pack(side=tk.LEFT, padx=(0, 20))
        
        self.cookie_status_var = tk.StringVar(value="未检测到Cookie")
        ttk.Label(row1, textvariable=self.cookie_status_var, foreground="gray").pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(row1, text="清除Cookie", command=self._clear_cookies, width=10).pack(side=tk.LEFT)
        
        self._check_cookie_status()
        
        # === 日志设置 ===
        log_frame = ttk.LabelFrame(parent, text="日志设置", padding="10")
        log_frame.pack(fill=tk.X, pady=(0, 10))
        
        row2 = ttk.Frame(log_frame)
        row2.pack(fill=tk.X, pady=2)
        
        self.log_to_file_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(row2, text="保存日志到文件", variable=self.log_to_file_var).pack(side=tk.LEFT, padx=(0, 20))
        
        ttk.Button(row2, text="打开日志文件", command=self._open_log_file).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(row2, text="清空日志", command=self._clear_log_file).pack(side=tk.LEFT)
        
        # === 速度控制 ===
        speed_frame = ttk.LabelFrame(parent, text="速度控制", padding="10")
        speed_frame.pack(fill=tk.X, pady=(0, 10))
        
        row3 = ttk.Frame(speed_frame)
        row3.pack(fill=tk.X, pady=2)
        
        ttk.Label(row3, text="点击延迟(秒):").pack(side=tk.LEFT)
        self.click_min_var = tk.StringVar(value="0.3")
        ttk.Entry(row3, textvariable=self.click_min_var, width=5).pack(side=tk.LEFT, padx=2)
        ttk.Label(row3, text="-").pack(side=tk.LEFT)
        self.click_max_var = tk.StringVar(value="0.5")
        ttk.Entry(row3, textvariable=self.click_max_var, width=5).pack(side=tk.LEFT, padx=(2, 20))
        
        ttk.Label(row3, text="滚动延迟(秒):").pack(side=tk.LEFT)
        self.scroll_min_var = tk.StringVar(value="0.4")
        ttk.Entry(row3, textvariable=self.scroll_min_var, width=5).pack(side=tk.LEFT, padx=2)
        ttk.Label(row3, text="-").pack(side=tk.LEFT)
        self.scroll_max_var = tk.StringVar(value="0.6")
        ttk.Entry(row3, textvariable=self.scroll_max_var, width=5).pack(side=tk.LEFT)
        
        # === 反爬设置 ===
        anti_frame = ttk.LabelFrame(parent, text="反爬虫设置", padding="10")
        anti_frame.pack(fill=tk.X, pady=(0, 10))
        
        row4 = ttk.Frame(anti_frame)
        row4.pack(fill=tk.X, pady=2)
        
        self.random_delay_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(row4, text="随机延迟（模拟人类行为）", variable=self.random_delay_var).pack(side=tk.LEFT, padx=(0, 20))
        
        self.random_scroll_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(row4, text="随机滚动距离", variable=self.random_scroll_var).pack(side=tk.LEFT)
        
        # === 数据库设置 ===
        db_frame = ttk.LabelFrame(parent, text="数据库设置", padding="10")
        db_frame.pack(fill=tk.X, pady=(0, 10))
        
        row5 = ttk.Frame(db_frame)
        row5.pack(fill=tk.X, pady=2)
        
        ttk.Label(row5, text="数据库路径:").pack(side=tk.LEFT)
        self.db_path_var = tk.StringVar(value="data/redbook.db")
        ttk.Entry(row5, textvariable=self.db_path_var, width=40).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(row5, text="浏览", command=self._browse_db_path).pack(side=tk.LEFT)
    
    # === 事件处理 ===
    def _on_mode_change(self):
        """切换爬取模式"""
        mode = self.crawl_type_var.get()
        
        # 禁用/启用相应输入框
        self.keyword_entry.config(state=tk.NORMAL if mode == "keyword" else tk.DISABLED)
        self.blogger_entry.config(state=tk.NORMAL if mode == "blogger" else tk.DISABLED)
        self.hot_combo.config(state="readonly" if mode == "hot" else tk.DISABLED)
    
    def _check_cookie_status(self):
        """检查Cookie状态"""
        if self.cookie_mgr.exists():
            saved_time = self.cookie_mgr.get_saved_time()
            if saved_time and saved_time != '未知':
                try:
                    dt = datetime.fromisoformat(saved_time)
                    time_str = dt.strftime("%m-%d %H:%M")
                    self.cookie_status_var.set(f"[已保存] Cookie ({time_str})")
                except Exception:
                    self.cookie_status_var.set("[已保存] Cookie")
            else:
                self.cookie_status_var.set("[已保存] Cookie")
        else:
            self.cookie_status_var.set("[未保存] 未检测到Cookie")
    
    def _use_saved_cookies(self):
        """使用已保存的Cookie"""
        if self.cookie_mgr.exists():
            saved_time = self.cookie_mgr.get_saved_time()
            msg = "将在爬取时自动加载Cookie，可跳过登录"
            if saved_time and saved_time != '未知':
                msg += f"\n\n保存时间: {saved_time}"
            messagebox.showinfo("Cookie信息", msg)
        else:
            messagebox.showwarning("提示", "未找到保存的Cookie\n请先完成一次登录，系统会自动保存")
    
    def _clear_cookies(self):
        """清除已保存的Cookie"""
        if self.cookie_mgr.exists():
            if messagebox.askyesno("确认", "确定要清除已保存的Cookie吗？\n清除后下次需要重新登录"):
                self.cookie_mgr.clear()
                self._check_cookie_status()
                self.log("Cookie已清除", "INFO")
        else:
            messagebox.showinfo("提示", "没有保存的Cookie")
    
    # === 预设 ===
    def _preset_turbo(self):
        self.crawl_mode_var.set("turbo")
        self.download_images_var.set(True)
        self.get_all_images_var.set(False)
        self.download_videos_var.set(False)
        self.get_content_var.set(False)
        self.get_comments_var.set(False)
        self.log("已应用极速采集预设", "SUCCESS")
    
    def _preset_complete(self):
        self.crawl_mode_var.set("standard")
        self.download_images_var.set(True)
        self.get_all_images_var.set(True)
        self.download_videos_var.set(True)
        self.get_content_var.set(True)
        self.get_tags_var.set(True)
        self.get_comments_var.set(True)
        self.log("已应用完整数据预设", "SUCCESS")
    
    def _preset_images(self):
        self.download_images_var.set(True)
        self.get_all_images_var.set(True)
        self.download_videos_var.set(False)
        self.get_content_var.set(False)
        self.get_comments_var.set(False)
        self.log("已应用只下图片预设", "SUCCESS")
    
    def _preset_videos(self):
        self.download_images_var.set(False)
        self.download_videos_var.set(True)
        self.note_type_var.set("视频")
        self.log("已应用只下视频预设", "SUCCESS")
    
    def _preset_text(self):
        self.download_images_var.set(False)
        self.download_videos_var.set(False)
        self.get_content_var.set(True)
        self.get_tags_var.set(True)
        self.get_comments_var.set(True)
        self.log("已应用只要文本预设", "SUCCESS")
    
    # === 日志 ===
    def log(self, message, level="INFO"):
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_queue.put((f"[{timestamp}] {message}\n", level))
        
        if self.config.log_to_file:
            self.file_logger.log(message, level)
    
    def _start_log_consumer(self):
        def consume():
            try:
                while True:
                    msg, level = self.log_queue.get_nowait()
                    self.log_text.config(state=tk.NORMAL)
                    self.log_text.insert(tk.END, msg, level)
                    self.log_text.see(tk.END)
                    self.log_text.config(state=tk.DISABLED)
            except queue.Empty:
                pass
            self.root.after(100, consume)
        self.root.after(100, consume)
    
    def _update_ui(self, **kwargs):
        if "status" in kwargs:
            self.status_var.set(kwargs["status"])
        if "notes" in kwargs:
            self.notes_var.set(kwargs["notes"])
        if "images" in kwargs:
            self.images_var.set(kwargs["images"])
        if "videos" in kwargs:
            self.videos_var.set(kwargs["videos"])
        if "time" in kwargs:
            self.time_var.set(kwargs["time"])
        if "progress" in kwargs:
            self.total_progress["value"] = kwargs["progress"]
            self.progress_label.config(text=f"{int(kwargs['progress'])}%")
    
    def _update_dashboard(self, stats):
        for key, value in stats.items():
            if key in self.dashboard_labels:
                self.dashboard_labels[key].config(text=str(int(value) if isinstance(value, float) else value))
    
    # === 爬取控制 ===
    def _start_crawl(self):
        if self.is_running:
            return
        
        # 检查输入
        crawl_type = self.crawl_type_var.get()
        if crawl_type == "keyword":
            # 空关键词表示爬取主页推荐
            pass
        elif crawl_type == "blogger":
            blogger_url = self.blogger_url_var.get().strip()
            if not blogger_url:
                messagebox.showwarning("提示", "请输入博主主页URL")
                return
        
        self._get_config()
        self._run_crawl()
    
    def _stop_crawl(self):
        self.should_stop = True
        self.log("正在停止...", "WARNING")
        self._update_ui(status="正在停止...")
        self.root.update()
    
    def _restore_gui_settings(self):
        """从配置恢复GUI设置"""
        try:
            # 基础设置
            self.keyword_var.set(self.config.keyword or "")
            self.scroll_var.set(str(self.config.scroll_times))
            self.max_notes_var.set(str(self.config.max_notes))
            self.parallel_var.set(str(self.config.parallel_downloads))
            self.crawl_mode_var.set(self.config.crawl_mode)
            self.crawl_type_var.set(self.config.crawl_type)
            self.blogger_url_var.set(self.config.blogger_url or "")
            
            # 筛选条件
            self.min_likes_var.set(str(self.config.min_likes))
            self.max_likes_var.set(str(self.config.max_likes))
            self.note_type_var.set(self.config.note_type_filter)
            
            # 内容选项
            self.get_content_var.set(self.config.get_content)
            self.get_tags_var.set(self.config.get_tags)
            self.get_time_var.set(self.config.get_publish_time)
            self.get_interactions_var.set(self.config.get_interactions)
            self.download_images_var.set(self.config.download_images)
            self.get_all_images_var.set(self.config.get_all_images)
            self.download_videos_var.set(self.config.download_videos)
            self.get_comments_var.set(self.config.get_comments)
            self.comments_count_var.set(str(self.config.comments_count))
            
            # 导出选项
            self.export_format_var.set(self.config.export_format)
            self.export_db_var.set(self.config.export_to_db)
            
            self.log("已恢复上次的设置", "SUCCESS")
        except Exception as e:
            self.log(f"恢复设置失败: {e}", "WARNING")
    
    def _save_gui_settings(self):
        """保存GUI设置到配置"""
        try:
            # 基础设置
            self.config.keyword = self.keyword_var.get().strip()
            self.config.scroll_times = int(self.scroll_var.get() or 10)
            self.config.max_notes = int(self.max_notes_var.get() or 30)
            self.config.parallel_downloads = int(self.parallel_var.get() or 10)
            self.config.crawl_mode = self.crawl_mode_var.get()
            self.config.crawl_type = self.crawl_type_var.get()
            self.config.blogger_url = self.blogger_url_var.get().strip()
            
            # 筛选条件
            self.config.min_likes = int(self.min_likes_var.get() or 0)
            self.config.max_likes = int(self.max_likes_var.get() or 999999)
            self.config.note_type_filter = self.note_type_var.get()
            
            # 内容选项
            self.config.get_content = self.get_content_var.get()
            self.config.get_tags = self.get_tags_var.get()
            self.config.get_publish_time = self.get_time_var.get()
            self.config.get_interactions = self.get_interactions_var.get()
            self.config.download_images = self.download_images_var.get()
            self.config.get_all_images = self.get_all_images_var.get()
            self.config.download_videos = self.download_videos_var.get()
            self.config.get_comments = self.get_comments_var.get()
            self.config.comments_count = int(self.comments_count_var.get() or 10)
            
            # 导出选项
            self.config.export_format = self.export_format_var.get()
            self.config.export_to_db = self.export_db_var.get()
        except Exception:
            pass
    
    def _get_config(self):
        """获取配置"""
        self.config.keyword = self.keyword_var.get().strip()
        self.config.crawl_type = self.crawl_type_var.get()
        self.config.blogger_url = self.blogger_url_var.get().strip()
        self.config.scroll_times = int(self.scroll_var.get())
        self.config.max_notes = int(self.max_notes_var.get())
        self.config.parallel_downloads = int(self.parallel_var.get())
        self.config.crawl_mode = self.crawl_mode_var.get()
        
        self.config.download_images = self.download_images_var.get()
        self.config.download_videos = self.download_videos_var.get()
        self.config.get_all_images = self.get_all_images_var.get()
        self.config.get_content = self.get_content_var.get()
        self.config.get_tags = self.get_tags_var.get()
        self.config.get_publish_time = self.get_time_var.get()
        self.config.get_comments = self.get_comments_var.get()
        self.config.comments_count = int(self.comments_count_var.get())
        self.config.get_interactions = self.get_interactions_var.get()
        
        self.config.min_likes = int(self.min_likes_var.get() or 0)
        self.config.max_likes = int(self.max_likes_var.get() or 999999)
        self.config.note_type_filter = self.note_type_var.get()
        
        self.config.export_format = self.export_format_var.get()
        self.config.export_to_db = self.export_db_var.get()
        self.config.save_cookies = self.save_cookies_var.get()
        self.config.log_to_file = self.log_to_file_var.get()
        
        self.config.click_delay = (float(self.click_min_var.get()), float(self.click_max_var.get()))
        self.config.scroll_delay = (float(self.scroll_min_var.get()), float(self.scroll_max_var.get()))
        
        self.downloader.max_workers = self.config.parallel_downloads
    
    def _run_crawl(self):
        self.is_running = True
        self.should_stop = False
        self.all_notes_data = []
        
        # 清空表格UI
        for item in self.result_tree.get_children():
            self.result_tree.delete(item)
        self.result_count_label.config(text="共 0 条记录")
        
        # 清空预览区域
        self.preview_canvas.delete("all")
        self.preview_images = []
        self.preview_image_paths = []
        self.current_video_path = None
        self.current_selected_note = None
        
        # 清空详情区域
        self.detail_text.config(state=tk.NORMAL)
        self.detail_text.delete(1.0, tk.END)
        self.detail_text.config(state=tk.DISABLED)
        
        # 清空批次数据
        self.batch_notes_data = []
        self.current_batch_folder = None
        
        # 确保数据源是"当前爬取"
        self.data_source_var.set("当前爬取")
        
        self.start_btn.config(state=tk.DISABLED)
        self.stop_btn.config(state=tk.NORMAL)
        
        thread = threading.Thread(target=self._crawl_thread, daemon=True)
        thread.start()
    
    def _crawl_thread(self):
        """爬取主线程（优化版，增强错误恢复）"""
        start_time = time.time()
        page = None
        total_notes = 0
        total_images = 0
        total_videos = 0
        error_count = 0
        MAX_ERRORS = 5  # 连续错误上限
        
        try:
            # 处理多关键词（空关键词表示爬取主页）
            keywords = [k.strip() for k in self.config.keyword.split(',') if k.strip()]
            if not keywords:
                keywords = [""]  # 空字符串表示主页
            
            for kw_idx, keyword in enumerate(keywords):
                if self.should_stop:
                    self.log("用户停止爬取", "WARNING")
                    break
                
                if error_count >= MAX_ERRORS:
                    self.log(f"连续错误超过{MAX_ERRORS}次，停止爬取", "ERROR")
                    break
                
                display_keyword = keyword if keyword else "主页推荐"
                self.log(f"开始爬取 [{kw_idx+1}/{len(keywords)}]: {display_keyword}", "INFO")
                
                # 复用浏览器实例（保持登录状态）
                if page is None:
                    if self.browser_page is not None:
                        # 复用已有的浏览器
                        page = self.browser_page
                        self.log("复用已打开的浏览器", "INFO")
                        
                        # 检查登录状态
                        page.get('https://www.xiaohongshu.com')
                        time.sleep(1.5)
                        if not self._check_login(page):
                            self.log("需要重新登录", "WARNING")
                            self._wait_for_login(page)
                    else:
                        # 首次启动浏览器
                        try:
                            user_data_dir = os.path.abspath("data/browser_profile")
                            os.makedirs(user_data_dir, exist_ok=True)
                            
                            co = ChromiumOptions()
                            co.set_user_data_path(user_data_dir)
                            co.set_argument('--no-first-run')
                            co.set_argument('--no-default-browser-check')
                            
                            page = ChromiumPage(co)
                            self.browser_page = page  # 保存实例以便复用
                            self.log("浏览器启动成功", "SUCCESS")
                        except Exception as e:
                            self.log(f"浏览器启动失败: {e}", "ERROR")
                            return
                        
                        # 访问小红书并检查登录状态
                        page.get('https://www.xiaohongshu.com')
                        time.sleep(2)
                        
                        if self._check_login(page):
                            self.log("登录状态有效", "SUCCESS")
                            # 获取Cookie传递给下载器
                            self._sync_browser_cookies(page)
                        else:
                            self.log("需要登录", "WARNING")
                            self._wait_for_login(page)
                            # 登录后获取Cookie
                            self._sync_browser_cookies(page)
                
                if self.should_stop:
                    break
                
                try:
                    # 访问页面（空关键词=主页，有关键词=搜索页）
                    if keyword:
                        keyword_code = quote(quote(keyword.encode('utf-8')).encode('gb2312'))
                        target_url = f'https://www.xiaohongshu.com/search_result?keyword={keyword_code}&source=web_search_result_notes'
                        self.log(f"访问搜索页面...", "INFO")
                        self._update_ui(status=f"搜索: {keyword}")
                    else:
                        target_url = 'https://www.xiaohongshu.com/explore'
                        self.log(f"访问主页推荐...", "INFO")
                        self._update_ui(status="爬取主页")
                    
                    page.get(target_url)
                    time.sleep(1.5)
                    
                    # 再次检查登录状态（搜索页可能弹出登录框）
                    if not self._check_login(page):
                        self.log("搜索页需要登录", "WARNING")
                        # 尝试关闭登录弹窗
                        try:
                            close_btn = page.ele('css:.close-icon, [class*="close"]', timeout=0.5)
                            if close_btn:
                                close_btn.click()
                                time.sleep(0.3)
                        except Exception:
                            pass
                        # 如果还是没登录，等待用户登录
                        if not self._check_login(page):
                            self._wait_for_login(page)
                            page.get(target_url)
                            time.sleep(1.5)
                    
                    # 自动滚动加载笔记（直到达到目标数量或无法加载更多）
                    prev_count = 0
                    no_change_count = 0
                    target_notes = self.config.max_notes
                    scroll_count = 0
                    max_scrolls = 100  # 最大滚动次数，防止无限循环
                    
                    self.log(f"自动加载笔记，目标: {target_notes} 个", "INFO")
                    
                    while scroll_count < max_scrolls:
                        if self.should_stop:
                            break
                        
                        scroll_count += 1
                        self._update_ui(status=f"加载中...")
                        
                        # 多种滚动方式组合
                        try:
                            # 方式1: 滚动到最后一个笔记
                            notes = page.eles("css:section.note-item")
                            if notes:
                                notes[-1].scroll.to_see()
                                time.sleep(0.3)
                            
                            # 方式2: 滚动整个页面
                            page.run_js("window.scrollBy(0, window.innerHeight)")
                            time.sleep(0.3)
                            
                            # 方式3: 滚动到页面底部
                            page.run_js("window.scrollTo(0, document.body.scrollHeight)")
                        except Exception:
                            page.scroll.to_bottom()
                        
                        # 等待内容加载
                        time.sleep(random.uniform(0.6, 1.0))
                        
                        # 检测当前笔记数量
                        curr_count = len(page.eles("css:section.note-item", timeout=0.5))
                        
                        if curr_count >= target_notes:
                            self.log(f"已加载足够笔记 ({curr_count}/{target_notes})", "SUCCESS")
                            break
                        
                        if curr_count == prev_count:
                            no_change_count += 1
                            if no_change_count >= 5:
                                self.log(f"加载完成，共 {curr_count} 个笔记 (页面无更多内容)", "INFO")
                                break
                        else:
                            no_change_count = 0
                            if scroll_count % 5 == 0:  # 每5次滚动输出一次进度
                                self.log(f"已加载 {curr_count} 个笔记...", "INFO")
                        
                        prev_count = curr_count
                    
                    if self.should_stop:
                        break
                    
                    # 回到顶部，确保排序从第一个笔记开始
                    page.scroll.to_top()
                    time.sleep(0.3)
                    
                    # 获取笔记列表
                    note_elements = page.eles("css:section.note-item")[:self.config.max_notes]
                    note_count = len(note_elements)
                    
                    if note_count == 0:
                        self.log(f"未找到笔记，跳过关键词: {keyword}", "WARNING")
                        error_count += 1
                        continue
                    
                    self.log(f"找到 {note_count} 个笔记", "SUCCESS")
                    error_count = 0  # 重置错误计数
                    
                    # 根据模式选择爬取方法
                    if self.config.crawl_mode == "turbo":
                        notes, imgs, vids = self._fast_crawl(page, note_elements, keyword, start_time)
                    else:
                        notes, imgs, vids = self._standard_crawl(page, note_elements, keyword, start_time)
                    
                    total_notes += notes
                    total_images += imgs
                    total_videos += vids
                    
                except Exception as e:
                    self.log(f"爬取关键词 '{keyword}' 时出错: {e}", "ERROR")
                    error_count += 1
                    continue
            
            # 保存数据
            if self.all_notes_data:
                try:
                    save_name = keywords[0] if keywords[0] else "主页推荐"
                    if len(keywords) > 1:
                        save_name = "多关键词"
                    filename = self._save_data(self.all_notes_data, save_name)
                    self.log(f"数据已保存: {filename}", "SUCCESS")
                    
                    # 更新仪表盘
                    df = pd.DataFrame(self.all_notes_data)
                    stats = DataAnalyzer.generate_stats(df)
                    self.root.after(0, lambda s=stats: self._update_dashboard(s))
                except Exception as e:
                    self.log(f"保存数据失败: {e}", "ERROR")
            
            # 保存Cookie
            if page and self.config.save_cookies:
                try:
                    if self.cookie_mgr.save(page):
                        self.log("Cookie已保存，下次可自动登录", "SUCCESS")
                        self.root.after(0, self._check_cookie_status)
                except Exception:
                    pass
            
            elapsed = int(time.time() - start_time)
            status = "已停止" if self.should_stop else "完成"
            self._update_ui(
                status=status,
                notes=f"笔记: {total_notes}",
                images=f"图片: {total_images}",
                videos=f"视频: {total_videos}",
                time=f"用时: {elapsed}秒",
                progress=100
            )
            
            # 显示下载统计
            dl_stats = self.downloader.get_stats()
            if dl_stats['success'] > 0:
                mb = dl_stats['bytes'] / (1024 * 1024)
                self.log(f"下载统计: 成功 {dl_stats['success']}, 失败 {dl_stats['failed']}, 总计 {mb:.1f}MB", "INFO")
            
            self.log(f"爬取{status}！笔记: {total_notes}, 图片: {total_images}, 视频: {total_videos}", "SUCCESS")
            self.root.after(0, self._refresh_history)
            
        except InterruptedError:
            self.log("爬取已取消", "WARNING")
        except Exception as e:
            self.log(f"严重错误: {str(e)}", "ERROR")
            import traceback
            self.file_logger.log(traceback.format_exc(), "ERROR")
        finally:
            # 不关闭浏览器，保持登录状态
            # 浏览器会在程序退出时关闭
            
            # 重置下载器状态
            self.downloader.close()
            self.downloader.reset_stats()
            
            self.is_running = False
            self.root.after(0, lambda: self.start_btn.config(state=tk.NORMAL))
            self.root.after(0, lambda: self.stop_btn.config(state=tk.DISABLED))
    
    def _sync_browser_cookies(self, page):
        """将浏览器Cookie同步到下载器"""
        try:
            cookies = page.cookies()
            if cookies:
                self.downloader.set_cookies(cookies)
                self.log(f"  已同步 {len(cookies)} 个Cookie到下载器", "INFO")
        except Exception as e:
            self.log(f"  同步Cookie失败: {e}", "WARNING")
    
    def _check_login(self, page) -> bool:
        """检查是否已登录（优先检测登录弹窗）"""
        try:
            # ===== 第一优先级：检查是否有登录弹窗（未登录标志）=====
            # 登录弹窗存在时，底层页面元素仍可能存在，所以必须先检查弹窗
            
            # 检查二维码登录弹窗
            qrcode = page.ele('xpath://img[contains(@src, "qrcode")]', timeout=0.3)
            if qrcode:
                return False
            
            # 检查"登录后查看搜索结果"按钮
            login_hint = page.ele('xpath://span[contains(text(), "登录后查看") or contains(text(), "扫码登录") or contains(text(), "手机号登录")]', timeout=0.3)
            if login_hint:
                return False
            
            # 检查登录弹窗的关闭按钮（登录弹窗特有的close-icon）
            close_icon = page.ele('css:.close-icon', timeout=0.2)
            if close_icon:
                # 如果有关闭按钮，检查附近是否有登录相关文字
                try:
                    parent = close_icon.parent()
                    if parent:
                        parent_text = parent.text or ""
                        if "登录" in parent_text or "扫码" in parent_text:
                            return False
                except Exception:
                    pass
            
            # 检查红色登录按钮
            login_btn = page.ele('css:.login-btn, button.login-btn', timeout=0.2)
            if login_btn:
                # 确认是侧边栏的登录按钮（未登录状态）
                btn_text = login_btn.text or ""
                if "登录" in btn_text:
                    return False
            
            # ===== 第二优先级：检查已登录标志 =====
            
            # 检查侧边栏"我"区域是否有用户主页链接
            user_profile = page.ele('css:.user.side-bar-component a[href*="/user/profile/"]', timeout=0.3)
            if user_profile:
                return True
            
            # 检查侧边栏是否有用户头像
            avatar = page.ele('css:.side-bar .reds-avatar', timeout=0.2)
            if avatar:
                return True
            
            # 检查侧边栏文本
            try:
                sidebar = page.ele('css:.side-bar', timeout=0.2)
                if sidebar:
                    text = sidebar.text or ""
                    # 已登录时有"发现、发布、通知、我"且没有"登录"按钮文字
                    if "我" in text and "发现" in text and "登录" not in text:
                        return True
                    # 未登录时有"登录"按钮
                    if "登录" in text:
                        return False
            except Exception:
                pass
            
            # 默认认为未登录（更安全，让用户确认）
            return False
            
        except Exception:
            return False
    
    def _wait_for_login(self, page):
        """等待登录"""
        self.log("请在浏览器中完成登录", "WARNING")
        self._update_ui(status="等待登录...")
        
        login_event = threading.Event()
        cancelled = [False]
        
        def show_dialog():
            result = messagebox.askokcancel(
                "等待登录",
                "请在浏览器中完成登录\n\n登录完成后点击【确定】\n点击【取消】停止爬取"
            )
            if not result:
                cancelled[0] = True
                self.should_stop = True
            login_event.set()
        
        self.root.after(0, show_dialog)
        login_event.wait()
        
        if cancelled[0]:
            raise InterruptedError("用户取消")
        
        # 登录完成后立即保存Cookie
        if self.config.save_cookies:
            try:
                time.sleep(1)  # 等待Cookie完全写入
                if self.cookie_mgr.save(page):
                    self.log("Cookie已保存，下次可自动登录", "SUCCESS")
                    self.root.after(0, self._check_cookie_status)
            except Exception as e:
                self.log(f"Cookie保存失败: {e}", "WARNING")
    
    def _get_sorted_note_indices(self, page) -> List[int]:
        """获取按位置排序的笔记索引（从上到下、从左到右）
        
        按行分组排序：
        1. 先按top排序
        2. 识别行（top差距<80px的视为同一行）
        3. 每行内按left排序
        """
        try:
            script = """
            return (() => {
                const notes = document.querySelectorAll('section.note-item');
                if (notes.length === 0) return [];
                
                const positions = [];
                notes.forEach((n, i) => {
                    const rect = n.getBoundingClientRect();
                    positions.push({
                        domIndex: i,
                        left: Math.round(rect.left),
                        top: Math.round(rect.top)
                    });
                });
                
                // 按行分组排序
                positions.sort((a, b) => a.top - b.top);
                
                const rows = [];
                let currentRow = [positions[0]];
                let rowTop = positions[0].top;
                
                for (let i = 1; i < positions.length; i++) {
                    if (positions[i].top - rowTop < 80) {
                        currentRow.push(positions[i]);
                    } else {
                        rows.push(currentRow);
                        currentRow = [positions[i]];
                        rowTop = positions[i].top;
                    }
                }
                rows.push(currentRow);
                
                // 每行按left排序，合并结果
                const result = [];
                rows.forEach(row => {
                    row.sort((a, b) => a.left - b.left);
                    row.forEach(p => result.push(p.domIndex));
                });
                
                return result;
            })()
            """
            result = page.run_js(script)
            if result and isinstance(result, list):
                self.log(f"[排序] 结果: {result[:10]}...", "DEBUG") if len(result) > 10 else None
                return result
        except Exception as e:
            self.log(f"[排序] 失败: {e}", "WARNING")
        # 失败时返回默认顺序
        return list(range(len(page.eles("css:section.note-item", timeout=0.5))))
    
    def _standard_crawl(self, page, note_elements, keyword: str, start_time: float) -> Tuple[int, int, int]:
        """标准模式爬取（按DOM顺序，稳定可靠）"""
        success = 0
        images = 0
        videos = 0
        from datetime import datetime
        timestamp = int(time.time())
        # 每次爬取创建独立文件夹（关键词_日期_时间）
        time_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        folder_name = keyword if keyword else "主页推荐"
        images_dir = f"images/{folder_name}_{time_str}"
        self.current_crawl_dir = images_dir  # 保存当前爬取目录
        consecutive_fails = 0
        MAX_CONSECUTIVE_FAILS = 3
        
        # 已爬取的笔记URL去重
        crawled_urls = set()
        
        # 保存页面URL用于恢复
        if keyword:
            keyword_code = quote(quote(keyword.encode('utf-8')).encode('gb2312'))
            base_url = f'https://www.xiaohongshu.com/search_result?keyword={keyword_code}&source=web_search_result_notes'
        else:
            base_url = 'https://www.xiaohongshu.com/explore'
        
        # 按顺序爬取（每次从头遍历找未爬取的笔记，更稳定）
        target_notes = self.config.max_notes
        self.log(f"开始爬取，目标 {target_notes} 个笔记", "INFO")
        
        max_attempts = target_notes * 3  # 最大尝试次数
        attempt = 0
        
        while success < target_notes and attempt < max_attempts:
            if self.should_stop:
                break
            
            attempt += 1
            elapsed = int(time.time() - start_time)
            progress = (success / target_notes) * 100 if target_notes > 0 else 0
            self._update_ui(
                status=f"爬取 {success}/{target_notes}",
                notes=f"笔记: {success}",
                images=f"图片: {images}",
                videos=f"视频: {videos}",
                time=f"用时: {elapsed}秒",
                progress=progress
            )
            
            # 连续失败时重新加载页面
            if consecutive_fails >= MAX_CONSECUTIVE_FAILS:
                self.log("连续失败，重新加载页面", "WARNING")
                try:
                    page.get(base_url)
                    time.sleep(2)
                    # 滚动加载
                    for _ in range(5):
                        page.scroll.to_bottom()
                        time.sleep(0.5)
                except Exception:
                    break
                consecutive_fails = 0
            
            try:
                # 确保在目标页面
                current_url = page.url or ""
                if '/explore/' in current_url and 'xsec_token' in current_url:
                    # 在笔记详情弹窗页，返回
                    try:
                        page.run_js("history.back()")
                        time.sleep(0.5)
                    except Exception:
                        pass
                
                # 获取所有笔记元素
                elements = page.eles("css:section.note-item", timeout=1)
                if not elements:
                    self.log("未找到笔记元素，尝试滚动加载", "WARNING")
                    page.scroll.to_bottom()
                    time.sleep(1)
                    consecutive_fails += 1
                    continue
                
                # 从头遍历，找到第一个未爬取的笔记
                found_note = False
                for i, elem in enumerate(elements):
                    # 获取封面链接
                    cover_link = elem.ele('css:a.cover', timeout=0.1)
                    if not cover_link:
                        continue  # 跳过没有封面的（推荐卡片）
                    
                    # 检测推荐搜索卡片
                    if self._is_search_recommend_card(elem):
                        continue
                    
                    # 获取笔记URL并提取笔记ID用于去重（去掉token等变化的参数）
                    note_href = cover_link.attr('href') or ""
                    # 提取笔记ID（格式如 /explore/67fa8d98000000001d0052a6）
                    note_id = ""
                    if '/explore/' in note_href:
                        try:
                            note_id = note_href.split('/explore/')[1].split('?')[0]
                        except:
                            note_id = note_href
                    else:
                        note_id = note_href
                    
                    if note_id in crawled_urls:
                        continue  # 已爬取过
                    
                    # 找到了未爬取的笔记
                    found_note = True
                    
                    # 获取卡片标题
                    try:
                        card_title = elem.ele('css:.title, .note-title', timeout=0.1)
                        card_title_text = (card_title.text if card_title else "")[:20]
                    except Exception:
                        card_title_text = ""
                    
                    self.log(f"[{success+1}/{target_notes}] 位置{i+1}, 标题={card_title_text}", "INFO")
                    
                    # 点击笔记打开弹窗
                    elem.scroll.to_see()
                    time.sleep(0.1)
                    cover_link.click()
                    
                    time.sleep(random.uniform(*self.config.click_delay))
                    
                    # 等待弹窗内容加载
                    popup_loaded = False
                    for _ in range(10):
                        try:
                            if page.ele('css:.note-content, .note-text, .author-wrapper', timeout=0.1):
                                popup_loaded = True
                                break
                        except Exception:
                            pass
                        time.sleep(0.2)
                    
                    # 额外等待互动数据和图片轮播加载
                    if popup_loaded:
                        # 等待互动数据
                        for _ in range(5):
                            try:
                                if page.ele('css:.like-wrapper .count, .engage-bar .count', timeout=0.1):
                                    break
                            except Exception:
                                pass
                            time.sleep(0.2)
                        
                        # 等待图片轮播加载（关键！）
                        for _ in range(5):
                            try:
                                if page.ele('css:.swiper-slide img, .carousel img, [class*="slider"] img', timeout=0.2):
                                    break
                            except Exception:
                                pass
                            time.sleep(0.3)
                    
                    # 检查是否无法浏览
                    try:
                        unavailable = page.ele('xpath://div[contains(text(), "暂时无法浏览")]', timeout=0.2)
                        if unavailable:
                            self.log("笔记无法浏览，跳过", "WARNING")
                            crawled_urls.add(note_id)
                            page.run_js("history.back()")
                            time.sleep(0.3)
                            break  # 退出内层循环，继续外层循环
                    except Exception:
                        pass
                    
                    # 确保URL已更新（验证当前笔记）
                    current_url = page.url
                    if note_id and note_id not in current_url:
                        self.log(f"  URL未更新，等待跳转...", "DEBUG")
                        for _ in range(10):
                            time.sleep(0.3)
                            current_url = page.url
                            if note_id in current_url:
                                break
                    
                    # 提取数据
                    time.sleep(0.5)  # 增加等待时间确保图片加载
                    note_data = self._extract_full_note(page, success, images_dir, timestamp, keyword)
                    crawled_urls.add(note_id)
                    
                    if note_data and note_data.get('title'):
                        self.all_notes_data.append(note_data)
                        success += 1
                        images += note_data.get('image_count', 0)
                        videos += 1 if note_data.get('video_url') else 0
                        consecutive_fails = 0
                        
                        self.root.after(0, lambda d=note_data, n=success: self._add_result_to_table(d, n-1))
                        
                        if self.config.export_to_db:
                            self.db_mgr.insert_note(note_data)
                        
                        title = note_data.get('title', '')[:25]
                        likes = note_data.get('like_count', 0)
                        self.log(f"[{success}] {title}... ❤️{likes}", "SUCCESS")
                    else:
                        consecutive_fails += 1
                    
                    # 返回列表页
                    try:
                        page.run_js("history.back()")
                        time.sleep(0.4)
                    except Exception:
                        page.actions.key_down('Escape').key_up('Escape')
                        time.sleep(0.3)
                    
                    break  # 成功处理一个笔记，退出内层循环
                
                # 如果没找到未爬取的笔记，尝试滚动加载更多
                if not found_note:
                    prev_count = len(elements)
                    self.log(f"当前页面 {prev_count} 个笔记已全部处理，尝试加载更多...", "INFO")
                    
                    # 多次滚动尝试加载更多
                    loaded_more = False
                    for scroll_try in range(3):
                        page.scroll.to_bottom()
                        time.sleep(1)
                        new_elements = page.eles("css:section.note-item", timeout=0.5)
                        if len(new_elements) > prev_count:
                            self.log(f"加载了 {len(new_elements) - prev_count} 个新笔记", "INFO")
                            loaded_more = True
                            break
                    
                    if not loaded_more:
                        self.log(f"页面无法加载更多笔记，共爬取 {success} 个", "WARNING")
                        break  # 退出while循环
                
            except Exception as e:
                consecutive_fails += 1
                error_msg = str(e)[:50] if str(e) else "未知错误"
                self.log(f"爬取失败: {error_msg}", "ERROR")
                
                # 尝试返回列表页
                try:
                    page.run_js("history.back()")
                    time.sleep(0.5)
                except Exception:
                    pass
        
        self.log(f"爬取完成：成功 {success} 个笔记", "SUCCESS")
        return success, images, videos
    
    def _fast_crawl(self, page, note_elements, keyword, start_time):
        """极速模式爬取"""
        from datetime import datetime
        records = []
        timestamp = int(time.time())
        # 每次爬取创建独立文件夹（关键词_日期_时间）
        time_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        folder_name = keyword if keyword else "主页推荐"
        images_dir = f"images/{folder_name}_{time_str}"
        self.current_crawl_dir = images_dir  # 保存当前爬取目录
        total = len(note_elements)
        
        download_tasks = []
        
        for idx in range(total):
            if self.should_stop:
                break
            
            self._update_ui(
                status=f"扫描 {idx+1}/{total}",
                progress=(idx / total) * 50
            )
            
            try:
                elements = page.eles("css:section.note-item")
                if idx >= len(elements):
                    continue
                
                elem = elements[idx]
                
                title = ""
                try:
                    t = elem.ele('xpath:.//span[contains(@class, "title")]', timeout=0.2)
                    if t:
                        title = t.text or ""
                except:
                    pass
                
                if not title:
                    try:
                        lines = (elem.text or "").split('\n')
                        title = next((l for l in lines if 5 < len(l) < 100), f"笔记{idx+1}")
                    except:
                        title = f"笔记{idx+1}"
                
                author = ""
                try:
                    a = elem.ele('xpath:.//span[contains(@class, "name")]', timeout=0.2)
                    if a:
                        author = a.text or ""
                except:
                    pass
                
                img_url = ""
                try:
                    img = elem.ele('xpath:.//img', timeout=0.2)
                    if img:
                        img_url = img.attr('src') or ""
                except:
                    pass
                
                note_link = ""
                try:
                    link = elem.ele('xpath:.//a[contains(@href, "/explore/")]', timeout=0.2)
                    if link:
                        href = link.attr('href') or ""
                        note_link = 'https://www.xiaohongshu.com' + href if href.startswith('/') else href
                except:
                    pass
                
                record = {
                    'title': title[:100],
                    'author': author or "未知",
                    'note_link': note_link,
                    'note_type': '图文',
                    'keyword': keyword,
                    'image_urls': [img_url] if img_url else [],
                    'image_count': 1 if img_url else 0,
                    'batch_dir': images_dir,  # 保存批次目录
                }
                
                if img_url and self.config.download_images:
                    # 过滤表情包
                    if not self._is_emoji_image(img_url):
                        folder = f"{images_dir}/note_{idx+1}_{timestamp}"
                        ext = '.webp' if '.webp' in img_url else '.jpg'
                        path = f"{folder}/img_1{ext}"
                        download_tasks.append((img_url, path, len(records)))
                
                records.append(record)
                
            except:
                continue
        
        # 批量下载
        if download_tasks and self.config.download_images:
            self.log(f"下载 {len(download_tasks)} 张图片...", "INFO")
            
            def prog(done, total):
                self._update_ui(status=f"下载 {done}/{total}", progress=50 + (done/total)*50)
            
            results = self.downloader.download_batch(
                [(u, p) for u, p, _ in download_tasks],
                prog,
                lambda: self.should_stop
            )
            
            for url, path, rec_idx in download_tasks:
                if results.get(url):
                    # 存储绝对路径
                    abs_path = os.path.abspath(results[url])
                    records[rec_idx]['local_images'] = [abs_path]
        
        self.all_notes_data.extend(records)
        
        img_count = sum(1 for r in records if r.get('local_images'))
        return len(records), img_count, 0
    
    def _extract_full_note(self, page, idx: int, images_dir: str, timestamp: int, keyword: str) -> Optional[Dict]:
        """提取完整笔记数据（基于实际页面结构优化）"""
        try:
            # 调试：显示当前弹窗URL
            current_url = page.url or ""
            self.log(f"[DEBUG] 提取笔记 idx={idx}, URL={current_url[:80]}", "INFO")
            
            data = {'keyword': keyword, 'image_count': 0, 'batch_dir': images_dir}
            
            FAST_TIMEOUT = 0.2
            
            # 标题 - 从当前弹窗URL获取note_id，然后精确获取当前笔记的标题
            title = ""
            
            # 从URL获取当前笔记ID
            url_note_id = None
            if '/explore/' in current_url:
                url_note_id = current_url.split('/explore/')[-1].split('?')[0].split('/')[0]
            
            # 方法1: 使用JavaScript直接获取当前弹窗的标题（最可靠）
            try:
                js_title = page.run_js("""
                    return (() => {
                        // 优先从弹窗容器获取标题
                        const modal = document.querySelector('.note-detail-mask, [class*="noteContainer"], .note-container');
                        if (modal) {
                            const titleEl = modal.querySelector('.title, [class*="title"]');
                            if (titleEl && titleEl.textContent.trim().length > 2) {
                                return titleEl.textContent.trim();
                            }
                        }
                        
                        // 从 __INITIAL_STATE__ 获取当前笔记的标题
                        try {
                            const state = window.__INITIAL_STATE__;
                            if (state && state.note) {
                                // 从URL获取当前笔记ID
                                const urlMatch = window.location.href.match(/explore\\/([a-zA-Z0-9]+)/);
                                const noteId = urlMatch ? urlMatch[1] : state.note.currentNoteId;
                                
                                if (noteId && state.note.noteDetailMap && state.note.noteDetailMap[noteId]) {
                                    const noteData = state.note.noteDetailMap[noteId];
                                    if (noteData.note && noteData.note.title) {
                                        return noteData.note.title;
                                    }
                                }
                            }
                        } catch(e) {}
                        
                        return '';
                    })()
                """)
                if js_title and len(js_title.strip()) > 2:
                    title = js_title.strip()
                    self.log(f"[DEBUG] JS获取标题: {title[:30]}", "INFO")
            except Exception as e:
                self.log(f"[DEBUG] JS获取标题失败: {e}", "WARNING")
            
            # 方法2: CSS选择器备用（更精确的选择器）
            if not title:
                title_selectors = [
                    'css:.note-detail-mask .title',       # 弹窗内的标题
                    'css:[class*="noteContainer"] .title',
                    'css:.note-content .title',           # 图文笔记标题
                    'css:#detail-title',                  # 旧版选择器
                ]
                for sel in title_selectors:
                    try:
                        e = page.ele(sel, timeout=FAST_TIMEOUT)
                        if e and e.text and len(e.text.strip()) > 2:
                            title = e.text.strip()
                            self.log(f"[DEBUG] CSS找到标题: {title[:30]}", "INFO")
                            break
                    except Exception:
                        continue
            
            # 方法3: 如果没有标题（视频笔记），用内容第一行作为标题
            if not title:
                try:
                    content_el = page.ele('css:.note-detail-mask .note-text, [class*="noteContainer"] .note-text, .note-text', timeout=FAST_TIMEOUT)
                    if content_el and content_el.text:
                        first_line = content_el.text.strip().split('\n')[0]
                        if len(first_line) > 2:
                            title = first_line[:50]
                            self.log(f"[DEBUG] 视频笔记，用内容作标题: {title[:30]}", "INFO")
                except Exception:
                    pass
            
            data['title'] = title[:200] if title else f"笔记{idx+1}"
            
            # 作者 - 优先从弹窗内获取
            author = ""
            
            # 方法1: JavaScript从弹窗或__INITIAL_STATE__获取
            try:
                js_author = page.run_js("""
                    return (() => {
                        // 从弹窗内获取
                        const modal = document.querySelector('.note-detail-mask, [class*="noteContainer"], .note-container');
                        if (modal) {
                            const authorEl = modal.querySelector('.username, .author-wrapper .name, .user-info .name');
                            if (authorEl && authorEl.textContent.trim().length > 0 && authorEl.textContent.trim().length < 50) {
                                return authorEl.textContent.trim();
                            }
                        }
                        
                        // 从 __INITIAL_STATE__ 获取
                        try {
                            const state = window.__INITIAL_STATE__;
                            if (state && state.note) {
                                const urlMatch = window.location.href.match(/explore\\/([a-zA-Z0-9]+)/);
                                const noteId = urlMatch ? urlMatch[1] : state.note.currentNoteId;
                                if (noteId && state.note.noteDetailMap && state.note.noteDetailMap[noteId]) {
                                    const noteData = state.note.noteDetailMap[noteId];
                                    if (noteData.note && noteData.note.user && noteData.note.user.nickname) {
                                        return noteData.note.user.nickname;
                                    }
                                }
                            }
                        } catch(e) {}
                        return '';
                    })()
                """)
                if js_author and len(js_author.strip()) > 0:
                    author = js_author.strip()
            except Exception:
                pass
            
            # 方法2: CSS选择器备用（更精确）
            if not author:
                author_selectors = [
                    'css:.note-detail-mask .username',
                    'css:[class*="noteContainer"] .username',
                    'css:.author-wrapper .username',
                    'css:.author-wrapper .name',
                    'css:.user-info .name',
                ]
                for sel in author_selectors:
                    try:
                        e = page.ele(sel, timeout=FAST_TIMEOUT)
                        if e and e.text:
                            txt = e.text.strip()
                            if txt and len(txt) < 50:
                                author = txt
                                break
                    except Exception:
                        continue
            data['author'] = author or "未知"
            
            # 正文内容 - 优先从弹窗内获取
            if self.config.get_content:
                content = ""
                content_selectors = [
                    'css:.note-detail-mask .note-text',
                    'css:[class*="noteContainer"] .note-text',
                    'css:.note-text',
                    'css:.desc',
                    'css:#detail-desc',
                ]
                for sel in content_selectors:
                    try:
                        e = page.ele(sel, timeout=FAST_TIMEOUT)
                        if e and e.text:
                            txt = e.text.strip()
                            if len(txt) > len(content):  # 取最长的内容
                                content = txt
                    except Exception:
                        continue
                if content:
                    self.log(f"[DEBUG] 找到内容: {content[:50]}...", "INFO")
                data['content'] = content
                
                # 提取标签
                if self.config.get_tags and content:
                    tags = re.findall(r'#([^\s#]+)', content)
                    data['tags'] = list(set(tags))[:20]
            
            # 发布时间和IP地区 - 使用.date (格式如 "01-24 江西")
            if self.config.get_publish_time:
                pub_time = ""
                ip_region = ""
                try:
                    e = page.ele('css:.date', timeout=FAST_TIMEOUT)
                    if e:
                        full_text = (e.text or "").strip()
                        # 分离时间和地区
                        if " " in full_text:
                            parts = full_text.split(" ", 1)
                            pub_time = parts[0]
                            ip_region = parts[1] if len(parts) > 1 else ""
                        else:
                            pub_time = full_text
                except Exception:
                    pass
                data['publish_time'] = pub_time
                data['ip_region'] = ip_region
            
            # 互动数据 - 从当前弹窗获取（使用URL中的noteId确保准确）
            if self.config.get_interactions:
                data['like_count'] = 0
                data['collect_count'] = 0
                data['comment_count'] = 0
                try:
                    # 方法1: 从__INITIAL_STATE__获取当前笔记的互动数据（最可靠）
                    try:
                        interact_result = page.run_js("""
                            return (() => {
                                const parseNum = (text) => {
                                    if (!text) return 0;
                                    text = String(text).trim().toLowerCase();
                                    if (text.includes('万')) return Math.floor(parseFloat(text.replace('万', '')) * 10000);
                                    if (text.includes('k')) return Math.floor(parseFloat(text.replace('k', '')) * 1000);
                                    const num = parseInt(text.replace(/[^0-9]/g, ''));
                                    return isNaN(num) ? 0 : num;
                                };
                                
                                // 方法1: 从__INITIAL_STATE__获取（使用URL中的noteId）
                                try {
                                    const state = window.__INITIAL_STATE__;
                                    if (state && state.note) {
                                        const urlMatch = window.location.href.match(/explore\\/([a-zA-Z0-9]+)/);
                                        const noteId = urlMatch ? urlMatch[1] : state.note.currentNoteId;
                                        
                                        if (noteId && state.note.noteDetailMap && state.note.noteDetailMap[noteId]) {
                                            const noteData = state.note.noteDetailMap[noteId].note;
                                            if (noteData && noteData.interactInfo) {
                                                return JSON.stringify({
                                                    likes: parseNum(noteData.interactInfo.likedCount),
                                                    collects: parseNum(noteData.interactInfo.collectedCount),
                                                    comments: parseNum(noteData.interactInfo.commentCount)
                                                });
                                            }
                                        }
                                    }
                                } catch(e) {}
                                
                                // 方法2: 从当前弹窗的DOM获取
                                const modal = document.querySelector('.note-detail-mask, [class*="noteContainer"], .note-container');
                                const searchRoot = modal || document;
                                const bar = searchRoot.querySelector('.buttons.engage-bar-style, .engage-bar, .interact-container');
                                
                                if (bar) {
                                    const likeEl = bar.querySelector('.like-wrapper .count');
                                    const collectEl = bar.querySelector('.collect-wrapper .count');
                                    const chatEl = bar.querySelector('.chat-wrapper .count');
                                    
                                    return JSON.stringify({
                                        likes: parseNum(likeEl?.textContent),
                                        collects: parseNum(collectEl?.textContent),
                                        comments: parseNum(chatEl?.textContent)
                                    });
                                }
                                
                                return '';
                            })()
                        """)
                        if interact_result:
                            import json
                            interact_data = json.loads(interact_result)
                            if interact_data.get('likes', 0) > 0:
                                data['like_count'] = int(interact_data['likes'])
                            if interact_data.get('collects', 0) > 0:
                                data['collect_count'] = int(interact_data['collects'])
                            if interact_data.get('comments', 0) > 0:
                                data['comment_count'] = int(interact_data['comments'])
                    except Exception as e:
                        self.log(f"  JS获取互动数据失败: {e}", "WARNING")
                    
                    # 方法2: CSS选择器备用（限定在弹窗内）
                    if data['like_count'] == 0:
                        like_selectors = [
                            'css:.note-detail-mask .like-wrapper .count',
                            'css:[class*="noteContainer"] .like-wrapper .count',
                            'css:.engage-bar-style .like-wrapper .count',
                        ]
                        for sel in like_selectors:
                            try:
                                e = page.ele(sel, timeout=0.3)
                                if e and e.text:
                                    num = self._parse_num(e.text)
                                    if num > 0:
                                        data['like_count'] = num
                                        break
                            except:
                                pass
                    
                    if data['collect_count'] == 0:
                        collect_selectors = [
                            'css:.note-detail-mask .collect-wrapper .count',
                            'css:[class*="noteContainer"] .collect-wrapper .count',
                        ]
                        for sel in collect_selectors:
                            try:
                                e = page.ele(sel, timeout=0.3)
                                if e and e.text:
                                    num = self._parse_num(e.text)
                                    if num > 0:
                                        data['collect_count'] = num
                                        break
                            except:
                                pass
                    
                    if data['comment_count'] == 0:
                        comment_selectors = [
                            'css:.note-detail-mask .chat-wrapper .count',
                            'css:[class*="noteContainer"] .chat-wrapper .count',
                        ]
                        for sel in comment_selectors:
                            try:
                                e = page.ele(sel, timeout=0.3)
                                if e and e.text:
                                    num = self._parse_num(e.text)
                                    if num > 0:
                                        data['comment_count'] = num
                                        break
                            except:
                                pass
                    
                    # 记录获取到的数据
                    if data['like_count'] > 0 or data['collect_count'] > 0:
                        self.log(f"  互动: ❤️{data['like_count']} ⭐{data['collect_count']} 💬{data['comment_count']}", "INFO")
                    
                except Exception as e:
                    self.log(f"  获取互动数据失败: {e}", "WARNING")
            
            # 链接和ID
            current_url = page.url
            data['note_link'] = current_url if '/explore/' in current_url else ""
            note_id = ""
            if '/explore/' in current_url:
                # 提取ID：/explore/xxxxx?token=xxx
                note_id = current_url.split('/explore/')[-1].split('?')[0]
            data['note_id'] = note_id
            
            # 检测笔记类型并获取视频URL
            note_type = "图文"
            video_url = ""
            try:
                # 多次尝试检测视频元素（视频可能延迟加载）
                v = None
                for _ in range(3):
                    v = page.ele('xpath://video', timeout=0.3)
                    if v:
                        break
                    time.sleep(0.2)
                
                if v:
                    note_type = "视频"
                    self.log(f"  检测到视频元素", "INFO")
                    
                    # 等待视频数据加载
                    time.sleep(0.5)
                    
                    # 尝试从JavaScript获取视频URL（最可靠的方法）
                    try:
                        script = """
                        return (() => {
                            // 方法1: 从 __INITIAL_STATE__ 获取当前笔记的视频（使用URL中的noteId）
                            try {
                                if (window.__INITIAL_STATE__ && window.__INITIAL_STATE__.note) {
                                    // 从URL获取当前笔记ID（更可靠）
                                    const urlMatch = window.location.href.match(/explore\\/([a-zA-Z0-9]+)/);
                                    const noteId = urlMatch ? urlMatch[1] : window.__INITIAL_STATE__.note.currentNoteId;
                                    const noteMap = window.__INITIAL_STATE__.note.noteDetailMap;
                                    
                                    if (noteMap && noteId && noteMap[noteId]) {
                                        const currentNote = noteMap[noteId];
                                        if (currentNote && currentNote.note && currentNote.note.video) {
                                            const video = currentNote.note.video;
                                            // 优先使用 consumer.originVideoKey
                                            if (video.consumer && video.consumer.originVideoKey) {
                                                return 'https://sns-video-bd.xhscdn.com/' + video.consumer.originVideoKey;
                                            }
                                            // 备用：直接的URL
                                            if (video.url && !video.url.startsWith('blob:')) {
                                                return video.url;
                                            }
                                            // 备用：media.stream
                                            if (video.media && video.media.stream && video.media.stream.h264) {
                                                const streams = video.media.stream.h264;
                                                if (streams.length > 0 && streams[0].masterUrl) {
                                                    return streams[0].masterUrl;
                                                }
                                            }
                                        }
                                    }
                                }
                            } catch(e) { console.log('State error:', e); }
                            
                            // 方法2: 从页面script标签中提取（寻找当前页面的视频数据）
                            const scripts = document.querySelectorAll('script');
                            for (let s of scripts) {
                                const text = s.textContent || '';
                                if (text.length < 100) continue;
                                
                                // 模式1: originVideoKey
                                let match = text.match(/"originVideoKey"\\s*:\\s*"([^"]+)"/);
                                if (match && match[1].length > 10) {
                                    return 'https://sns-video-bd.xhscdn.com/' + match[1];
                                }
                                
                                // 模式2: masterUrl
                                match = text.match(/"masterUrl"\\s*:\\s*"(https?:[^"]+)"/);
                                if (match) return match[1].replace(/\\\\/g, '');
                                
                                // 模式3: 直接的视频URL
                                match = text.match(/(https?:\\/\\/sns-video[^"'\\s]+\\.mp4[^"'\\s]*)/);
                                if (match) return match[1];
                                
                                // 模式4: xhscdn视频URL
                                match = text.match(/(https?:\\/\\/[^"'\\s]*xhscdn\\.com[^"'\\s]*\\/video[^"'\\s]*)/);
                                if (match) return match[1];
                            }
                            
                            // 方法3: 从video元素获取
                            const video = document.querySelector('video');
                            if (video) {
                                if (video.src && !video.src.startsWith('blob:')) return video.src;
                                const source = video.querySelector('source');
                                if (source && source.src && !source.src.startsWith('blob:')) return source.src;
                            }
                            
                            return '';
                        })()
                        """
                        result = page.run_js(script)
                        if result and not result.startswith('blob:') and len(result) > 20:
                            video_url = result
                            self.log(f"  视频URL获取成功: {video_url[:60]}...", "SUCCESS")
                    except Exception as e:
                        self.log(f"  JS获取视频URL失败: {e}", "WARNING")
                    
                    # 如果JS方法失败，尝试从video元素直接获取
                    if not video_url or video_url.startswith('blob:'):
                        try:
                            video_url = v.attr('src') or ""
                            if video_url and video_url.startswith('blob:'):
                                video_url = ""
                        except:
                            pass
                    
                    if not video_url:
                        self.log(f"  无法获取可下载的视频URL (可能是blob格式)", "WARNING")
                        
            except Exception as e:
                self.log(f"  视频检测异常: {e}", "WARNING")
                
            data['note_type'] = note_type
            data['video_url'] = video_url
            
            # 获取图片URL - 优先使用JavaScript从页面状态获取
            preview_images = []
            try:
                # 方法1: 从当前弹窗的DOM直接获取图片（最可靠）
                # 先等待图片加载
                time.sleep(0.5)
                
                try:
                    # 从当前URL获取note_id
                    current_url = page.url
                    url_note_id = None
                    if '/explore/' in current_url:
                        url_note_id = current_url.split('/explore/')[-1].split('?')[0].split('/')[0]
                    
                    js_images = page.run_js("""
                        return (() => {
                            const images = [];
                            
                            // 方法1: 从当前可见的弹窗/详情页获取图片
                            // 查找笔记详情弹窗
                            const noteModal = document.querySelector('.note-detail-mask, .note-container, [class*="noteContainer"], [class*="note-detail"]');
                            const searchRoot = noteModal || document.body;
                            
                            // 获取所有图片轮播中的图片
                            const carouselImgs = searchRoot.querySelectorAll('.swiper-slide img, .carousel img, [class*="slider"] img, [class*="carousel"] img');
                            for (let img of carouselImgs) {
                                const src = img.src || img.getAttribute('data-src') || '';
                                if (src.length > 50 && (src.includes('xhscdn') || src.includes('sns-')) && 
                                    !src.includes('avatar') && !src.includes('emoji') && !src.includes('icon')) {
                                    images.push(src);
                                }
                            }
                            
                            // 如果轮播没找到，获取所有大图
                            if (images.length === 0) {
                                const allImgs = searchRoot.querySelectorAll('img');
                                for (let img of allImgs) {
                                    const src = img.src || '';
                                    // 只获取内容图片（大于一定尺寸或特定域名）
                                    if (src.length > 80 && (src.includes('xhscdn') || src.includes('sns-img') || src.includes('sns-webpic'))) {
                                        if (!src.includes('avatar') && !src.includes('emoji') && !src.includes('icon') && !src.includes('loading')) {
                                            // 检查图片尺寸
                                            if (img.naturalWidth > 100 || img.width > 100) {
                                                images.push(src);
                                            } else if (img.naturalWidth === 0) {
                                                // 图片可能还没加载，也加入
                                                images.push(src);
                                            }
                                        }
                                    }
                                }
                            }
                            
                            // 方法2: 尝试从 __INITIAL_STATE__ 获取（作为补充）
                            if (images.length === 0) {
                                try {
                                    const state = window.__INITIAL_STATE__;
                                    if (state && state.note && state.note.noteDetailMap) {
                                        // 从URL获取当前笔记ID
                                        const urlMatch = window.location.href.match(/explore\\/([a-zA-Z0-9]+)/);
                                        const noteId = urlMatch ? urlMatch[1] : state.note.currentNoteId;
                                        
                                        if (noteId && state.note.noteDetailMap[noteId]) {
                                            const noteData = state.note.noteDetailMap[noteId];
                                            if (noteData.note && noteData.note.imageList) {
                                                for (let img of noteData.note.imageList) {
                                                    const url = img.urlDefault || img.url;
                                                    if (url) images.push(url);
                                                }
                                            }
                                        }
                                    }
                                } catch(e) {}
                            }
                            
                            return JSON.stringify([...new Set(images)].slice(0, 20));
                        })()
                    """)
                    if js_images:
                        import json
                        preview_images = json.loads(js_images)
                        self.log(f"  JS获取到 {len(preview_images)} 张图片", "INFO")
                except Exception as e:
                    self.log(f"  JS获取图片失败: {e}", "WARNING")
                
                # 方法2: CSS选择器备用 - 更精确的选择器
                def get_current_images():
                    urls = []
                    # 优先从弹窗内的轮播获取
                    selectors = [
                        'css:.note-detail-mask .swiper-slide img',
                        'css:.note-container .swiper-slide img',
                        'css:[class*="noteContainer"] img',
                        'css:.swiper-wrapper img',
                        'css:.note-slider-img img',
                        'css:.carousel-img img',
                    ]
                    for sel in selectors:
                        try:
                            imgs = page.eles(sel, timeout=0.2)
                            if imgs:
                                for img in imgs[:20]:
                                    src = img.attr('src') or ""
                                    if src and len(src) > 50:
                                        src_lower = src.lower()
                                        if 'avatar' not in src_lower and 'icon' not in src_lower and 'emoji' not in src_lower:
                                            if not self._is_emoji_image(src):
                                                if src not in urls:
                                                    urls.append(src)
                                if urls:  # 找到就停止
                                    break
                        except:
                            pass
                    return urls
                
                # 如果JS没获取到，使用CSS选择器
                if not preview_images:
                    preview_images = get_current_images()
                
                # 如果开启了获取全部图片，尝试切换轮播获取更多
                if self.config.get_all_images and note_type != "视频":
                    # 尝试多种方式切换轮播
                    max_clicks = 15  # 最多点击15次
                    for click_idx in range(max_clicks):
                        if self.should_stop:
                            break
                        
                        # 尝试点击下一张按钮
                        next_clicked = False
                        next_selectors = [
                            'css:.next-btn',
                            'css:.swiper-button-next',
                            'css:.carousel-next',
                            'css:[class*="next"]',
                            'xpath://div[contains(@class, "arrow") and contains(@class, "right")]',
                            'xpath://button[contains(@class, "next")]',
                        ]
                        
                        for sel in next_selectors:
                            try:
                                next_btn = page.ele(sel, timeout=0.2)
                                if next_btn:
                                    next_btn.click()
                                    next_clicked = True
                                    time.sleep(0.3)
                                    break
                            except:
                                pass
                        
                        # 如果没找到按钮，尝试用键盘右箭头
                        if not next_clicked:
                            try:
                                page.actions.key_down('RIGHT').key_up('RIGHT')
                                time.sleep(0.3)
                            except:
                                pass
                        
                        # 获取新图片
                        new_images = get_current_images()
                        old_count = len(preview_images)
                        for img in new_images:
                            if img not in preview_images:
                                preview_images.append(img)
                        
                        # 如果没有新图片，说明已经到最后一张
                        if len(preview_images) == old_count:
                            break
                    
                    if len(preview_images) > 1:
                        self.log(f"  轮播获取到 {len(preview_images)} 张图片", "INFO")
                
            except Exception as e:
                self.log(f"  获取图片异常: {e}", "WARNING")
            
            # 过滤重复和Live图（Live图只保留一张）
            filtered_images = self._filter_live_images(preview_images)
            data['image_urls'] = filtered_images[:20]  # 最多保存20张
            self.log(f"  共获取到 {len(data['image_urls'])} 张图片URL", "INFO")
            
            # 批量下载图片
            if self.config.download_images and data['image_urls']:
                folder = f"{images_dir}/note_{idx+1}_{timestamp}"
                tasks = []
                for i, url in enumerate(data['image_urls'], 1):
                    ext = '.webp' if '.webp' in url else '.jpg'
                    tasks.append((url, f"{folder}/img_{i}{ext}"))
                
                if tasks:
                    results = self.downloader.download_batch(tasks, None, lambda: self.should_stop)
                    # 存储绝对路径
                    data['local_images'] = [os.path.abspath(r) for r in results.values() if r]
                    data['image_count'] = len(data['local_images'])
                    self.log(f"  下载成功 {data['image_count']}/{len(tasks)} 张图片", "SUCCESS" if data['image_count'] > 0 else "WARNING")
            elif not data['image_urls']:
                self.log(f"  未获取到图片URL", "WARNING")
            
            # 下载视频
            if self.config.download_videos and video_url:
                self.log(f"  开始下载视频...", "INFO")
                folder = f"{images_dir}/note_{idx+1}_{timestamp}"
                os.makedirs(folder, exist_ok=True)
                video_path = f"{folder}/video.mp4"
                result = self.downloader.download_file(video_url, video_path, lambda: self.should_stop, min_size=10240)
                if result:
                    data['local_video'] = result
                    file_size = os.path.getsize(result) if os.path.exists(result) else 0
                    self.log(f"  视频下载成功: {file_size/1024/1024:.1f}MB", "SUCCESS")
                else:
                    self.log(f"  视频下载失败", "WARNING")
            
            # 评论爬取（优化版）
            if self.config.get_comments:
                comments = self._extract_comments(page)
                data['comments'] = comments
                if comments:
                    self.log(f"  获取到 {len(comments)} 条评论", "INFO")
                    
                    # 下载评论图片到单独的 comments 文件夹
                    comment_images_urls = []
                    for comment in comments:
                        if comment.get('images'):
                            comment_images_urls.extend(comment.get('images', []))
                    
                    if comment_images_urls and self.config.download_images:
                        # 使用与笔记图片相同的文件夹路径
                        note_save_folder = f"{images_dir}/note_{idx+1}_{timestamp}"
                        comments_dir = os.path.join(note_save_folder, 'comments')
                        os.makedirs(comments_dir, exist_ok=True)
                        
                        comment_img_count = 0
                        for i, img_url in enumerate(comment_images_urls[:20]):  # 最多20张评论图片
                            try:
                                ext = '.jpg'
                                if '.png' in img_url.lower():
                                    ext = '.png'
                                elif '.webp' in img_url.lower():
                                    ext = '.webp'
                                
                                filename = f"comment_img_{i+1}{ext}"
                                filepath = os.path.join(comments_dir, filename)
                                
                                if self.downloader.download_with_session(img_url, filepath, page):
                                    comment_img_count += 1
                            except Exception:
                                pass
                        
                        if comment_img_count > 0:
                            self.log(f"  评论图片: {comment_img_count}张 (保存到 comments 文件夹)", "INFO")
                            data['comment_images_count'] = comment_img_count
            
            return data
            
        except Exception as e:
            self.log(f"提取数据失败: {e}", "ERROR")
            return None
    
    def _extract_single_comment(self, item, existing_contents: set) -> Optional[Dict]:
        """提取单条评论的完整信息"""
        exclude_words = {'关注', '点赞', '收藏', '分享', '复制', '举报', '回复', '查看', '展开', '赞', '条评论', '说点什么', '取消', '发送'}
        
        try:
            # 获取评论者名字
            name_el = item.ele('css:.name, .user-name, .author-name, .nickname', timeout=0.1)
            name = (name_el.text if name_el else "").strip()
            
            # 获取评论内容
            content_el = item.ele('css:.content, .comment-content, .note-text', timeout=0.1)
            content = (content_el.text if content_el else "").strip()
            
            # 过滤无效评论
            if not content or len(content) <= 3 or len(content) >= 500:
                return None
            if content in existing_contents:
                return None
            if content in exclude_words or content.isdigit():
                return None
            
            # 获取时间
            time_el = item.ele('css:.date, .time, .info .date, .comment-time', timeout=0.1)
            time_text = (time_el.text if time_el else "").strip()
            
            # 获取IP地址/地区
            ip_text = ""
            try:
                # 尝试多种选择器获取IP/地区
                ip_el = item.ele('css:.ip, .location, .region, .area', timeout=0.1)
                if ip_el:
                    ip_text = ip_el.text.strip()
                else:
                    # 从时间文本中提取地区（如 "3天前 浙江"）
                    if time_text and " " in time_text:
                        parts = time_text.split()
                        if len(parts) >= 2:
                            # 检查最后一部分是否像地区名
                            last_part = parts[-1]
                            if not any(c in last_part for c in ['前', '天', '小时', '分钟', '秒', '月', '年']):
                                ip_text = last_part
                                time_text = " ".join(parts[:-1])
            except Exception:
                pass
            
            # 获取点赞数
            like_count = 0
            try:
                like_el = item.ele('css:.like-count, .likes, .like-num, .zan-count, [class*="like"] span', timeout=0.1)
                if like_el:
                    like_text = like_el.text.strip()
                    # 解析点赞数（可能是 "1.2万" 或 "1234"）
                    if like_text:
                        if '万' in like_text:
                            like_count = int(float(like_text.replace('万', '')) * 10000)
                        elif like_text.isdigit():
                            like_count = int(like_text)
            except Exception:
                pass
            
            # 检测评论中是否有图片
            has_image = False
            comment_images = []
            try:
                imgs = item.eles('css:img.comment-img, .comment-image img, .comment-pic img', timeout=0.1)
                if imgs:
                    has_image = True
                    for img in imgs[:3]:  # 最多3张图
                        src = img.attr('src') or ""
                        if src and 'avatar' not in src.lower() and len(src) > 30:
                            comment_images.append(src)
            except Exception:
                pass
            
            return {
                'author': name or "匿名用户",
                'content': content,
                'time': time_text,
                'ip': ip_text,
                'likes': like_count,
                'has_image': has_image,
                'images': comment_images
            }
        except Exception:
            return None
    
    def _extract_comments(self, page) -> List[Dict]:
        """提取评论内容（基于浏览器自动化分析的实际DOM结构）
        返回包含评论者、内容、时间、IP、点赞数、图片标记的字典列表
        """
        comments = []
        max_count = self.config.comments_count
        existing_contents = set()
        
        try:
            # 获取所有评论项
            comment_items = page.eles('css:.comment-item, .parent-comment, .comment-inner', timeout=0.5)
            
            for item in comment_items:
                if len(comments) >= max_count:
                    break
                
                comment = self._extract_single_comment(item, existing_contents)
                if comment:
                    comments.append(comment)
                    existing_contents.add(comment['content'])
            
            # 如果还没有足够的评论，尝试滚动评论区加载更多
            if len(comments) < max_count:
                try:
                    comments_container = page.ele('css:.comments-container, .comments-el, .note-scroller', timeout=0.3)
                    if comments_container:
                        comments_container.scroll.to_bottom()
                        time.sleep(0.3)
                        
                        # 再次获取新加载的评论
                        new_items = page.eles('css:.comment-item, .comment-inner', timeout=0.3)
                        for item in new_items:
                            if len(comments) >= max_count:
                                break
                            comment = self._extract_single_comment(item, existing_contents)
                            if comment:
                                comments.append(comment)
                                existing_contents.add(comment['content'])
                except Exception:
                    pass
                    
        except Exception:
            pass
        
        return comments
    
    def _filter_live_images(self, image_urls: list) -> list:
        """过滤Live图（动态图片），只保留一张静态版本
        
        Live图特征：
        1. URL中包含 'live' 关键字
        2. 同一张图片有静态和动态两个版本
        3. URL结构相似，只是路径或参数不同
        """
        import re
        
        if not image_urls:
            return []
        
        # 去重
        unique_urls = list(dict.fromkeys(image_urls))
        
        def extract_image_id(url):
            """提取图片的核心ID（去掉所有变体标记）"""
            # 移除查询参数
            base = url.split('?')[0]
            # 移除处理参数如 !nd_dft_wlteh_webp_3
            base = re.sub(r'![^/]+$', '', base)
            
            # 提取文件名部分
            filename = base.split('/')[-1]
            
            # 移除扩展名
            filename = re.sub(r'\.(jpg|jpeg|png|webp|gif|heic)$', '', filename, flags=re.IGNORECASE)
            
            # 移除live相关标记
            # 例如: spectrum/1040g0k031fat0rfh5g6g5p4sk5ohqo95i4stbh0_live.jpg -> spectrum/1040g0k031fat0rfh5g6g5p4sk5ohqo95i4stbh0
            filename = re.sub(r'_live\d*$', '', filename)
            filename = re.sub(r'-live\d*$', '', filename)
            
            # 提取核心ID（通常是长字符串）
            # 匹配类似 1040g0k031fat0rfh5g6g5p4sk5ohqo95i4stbh0 的ID
            id_match = re.search(r'([a-z0-9]{20,})', filename, re.IGNORECASE)
            if id_match:
                return id_match.group(1).lower()
            
            return filename.lower()
        
        def is_live_url(url):
            """判断是否是Live图URL"""
            url_lower = url.lower()
            return 'live' in url_lower or '/live/' in url_lower
        
        # 按图片ID分组
        url_groups = {}
        for url in unique_urls:
            img_id = extract_image_id(url)
            if img_id not in url_groups:
                url_groups[img_id] = []
            url_groups[img_id].append(url)
        
        # 每组只保留一张（优先非live的静态图）
        filtered = []
        for img_id, urls in url_groups.items():
            if len(urls) == 1:
                filtered.append(urls[0])
            else:
                # 多张相似图片，选择最优的一张
                # 优先级：不含live > 含jpg/png > 其他
                best = None
                for url in urls:
                    if not is_live_url(url):
                        # 优先选择静态图
                        url_lower = url.lower()
                        if best is None:
                            best = url
                        elif '.jpg' in url_lower or '.png' in url_lower:
                            best = url
                
                # 如果全是live图，取第一张
                if best is None:
                    best = urls[0]
                
                filtered.append(best)
                self.log(f"  Live图过滤: {len(urls)}张相似图 -> 保留1张", "DEBUG")
        
        return filtered
    
    def _is_emoji_image(self, url: str) -> bool:
        """检测是否是表情包图片"""
        if not url:
            return False
        url_lower = url.lower()
        
        import re
        
        # 1. URL关键词检测 - 扩展关键词列表
        emoji_keywords = [
            'emoji', 'sticker', 'emote', 'emoticon', 'expression',
            'spectrum', 'meme', 'gif', 'animated', 
            '/e/', '/em/', '/stk/', '/stick/'
        ]
        for kw in emoji_keywords:
            if kw in url_lower:
                return True
        
        # 2. 小红书表情包特征：通常是小尺寸图片
        # 检测URL中的尺寸参数，如 /w/120 或 imageView2/2/w/200 或 !nd_
        size_patterns = [
            r'/w/(\d+)',
            r'/h/(\d+)', 
            r'imageview2/\d/w/(\d+)',
            r'!nd_dft_wlteh_webp_(\d+)',
            r'_(\d+)x(\d+)\.',
        ]
        for pattern in size_patterns:
            match = re.search(pattern, url_lower)
            if match:
                try:
                    size = int(match.group(1))
                    if size <= 300:  # 宽度小于300像素，可能是表情
                        return True
                except:
                    pass
        
        # 3. 检测表情包CDN特征
        emoji_cdn_patterns = [
            'fe-static',
            '/emoji/',
            'spectrum.xhscdn',
            'sticker.xhscdn',
            'ci.xiaohongshu.com/spectrum',
        ]
        for pattern in emoji_cdn_patterns:
            if pattern in url_lower:
                return True
        
        # 4. 检测非常短的图片URL（通常是内联表情）
        if len(url) < 100:
            return True
        
        # 5. 检测URL中没有常规图片路径特征（正常笔记图片通常有特定路径）
        normal_patterns = ['sns-img', 'sns-webpic', 'note', 'traceId']
        has_normal_pattern = any(p in url_lower for p in normal_patterns)
        if not has_normal_pattern and 'xhscdn' in url_lower:
            # 小红书CDN但不是常规图片路径，可能是表情
            return True
            
        return False
    
    def _is_search_recommend_card(self, elem):
        """检测是否是'大家都在搜'推荐卡片"""
        try:
            # 获取卡片的文本内容
            text = elem.text or ""
            
            # 检测推荐搜索卡片的特征
            if "大家都在搜" in text:
                return True
            if "热门搜索" in text:
                return True
            
            # 检测卡片内是否有推荐搜索相关的class
            html = elem.html or ""
            if "search-recommend" in html.lower():
                return True
            if "hot-search" in html.lower():
                return True
            
            # 检测是否有多个搜索关键词链接（推荐卡片的特征）
            try:
                links = elem.eles('css:a')
                # 推荐卡片通常有多个链接，且没有封面图片
                cover = elem.ele('css:a.cover, .cover', timeout=0.1)
                if len(links) > 3 and not cover:
                    return True
            except:
                pass
                
        except Exception:
            pass
        return False
    
    def _parse_num(self, text) -> int:
        """解析数字（支持万/k单位）"""
        if not text:
            return 0
        text = str(text).strip().lower()
        try:
            if '万' in text:
                return int(float(text.replace('万', '')) * 10000)
            if 'k' in text:
                return int(float(text.replace('k', '')) * 1000)
            return int(re.sub(r'[^\d]', '', text) or 0)
        except Exception:
            return 0
    
    def _save_data(self, data, keyword):
        """保存数据"""
        os.makedirs("data", exist_ok=True)
        timestamp = int(time.time())
        
        # 预处理数据 - 将复杂类型转换为字符串
        processed_data = []
        for item in data:
            processed_item = item.copy()
            
            # 处理评论 - 将字典列表转为可读字符串（包含完整信息）
            if 'comments' in processed_item and isinstance(processed_item['comments'], list):
                comments = processed_item['comments']
                if comments and isinstance(comments[0], dict):
                    # 格式: "[用户名|IP|时间|赞数|图片标记] 评论内容"
                    comment_strs = []
                    for i, c in enumerate(comments, 1):
                        author = c.get('author', '') or '匿名'
                        content = c.get('content', '')
                        time_str = c.get('time', '')
                        ip_str = c.get('ip', '')
                        likes = c.get('likes', 0)
                        has_image = c.get('has_image', False)
                        
                        if content:
                            # 构建评论信息
                            info_parts = [f"@{author}"]
                            if ip_str:
                                info_parts.append(ip_str)
                            if time_str:
                                info_parts.append(time_str)
                            if likes > 0:
                                info_parts.append(f"❤️{likes}")
                            if has_image:
                                info_parts.append("[含图]")
                            
                            info = " | ".join(info_parts)
                            comment_strs.append(f"[{i}] {info}: {content}")
                    processed_item['comments'] = '\n'.join(comment_strs)
                else:
                    processed_item['comments'] = '\n'.join(str(c) for c in comments)
            
            # 处理标签列表
            if 'tags' in processed_item and isinstance(processed_item['tags'], list):
                processed_item['tags'] = ', '.join(processed_item['tags'])
            
            # 处理图片URL列表
            if 'image_urls' in processed_item and isinstance(processed_item['image_urls'], list):
                processed_item['image_urls'] = ' | '.join(processed_item['image_urls'])
            
            # 处理本地图片路径列表
            if 'local_images' in processed_item and isinstance(processed_item['local_images'], list):
                processed_item['local_images'] = ' | '.join(processed_item['local_images'])
            
            processed_data.append(processed_item)
        
        # 转换为DataFrame
        df = pd.DataFrame(processed_data)
        
        # 英文列名到中文列名的映射
        column_mapping = {
            'keyword': '搜索关键词',
            'title': '标题',
            'author': '作者',
            'content': '正文内容',
            'tags': '标签',
            'publish_time': '发布时间',
            'ip_region': 'IP地区',
            'like_count': '点赞数',
            'collect_count': '收藏数',
            'comment_count': '评论数',
            'comments': '评论内容',
            'note_type': '笔记类型',
            'note_link': '笔记链接',
            'note_id': '笔记ID',
            'video_url': '视频链接',
            'image_urls': '图片链接',
            'image_count': '图片数量',
            'local_images': '本地图片路径',
            'local_video': '本地视频路径',
        }
        
        # 重命名列
        df = df.rename(columns=column_mapping)
        
        ext = self.config.export_format
        filename = f"data/搜索结果_{keyword}_{timestamp}.{ext}"
        
        if ext == "xlsx":
            df.to_excel(filename, index=False)
        elif ext == "csv":
            df.to_csv(filename, index=False, encoding='utf-8-sig')
        elif ext == "json":
            # JSON格式保留原始结构
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        
        # 同时保存一份到当前爬取目录
        if hasattr(self, 'current_crawl_dir') and self.current_crawl_dir:
            try:
                os.makedirs(self.current_crawl_dir, exist_ok=True)
                crawl_file = f"{self.current_crawl_dir}/搜索结果.{ext}"
                if ext == "xlsx":
                    df.to_excel(crawl_file, index=False)
                elif ext == "csv":
                    df.to_csv(crawl_file, index=False, encoding='utf-8-sig')
                elif ext == "json":
                    with open(crawl_file, 'w', encoding='utf-8') as f:
                        json.dump(data, f, ensure_ascii=False, indent=2)
            except Exception:
                pass
        
        return filename
    
    # === 分析功能 ===
    def _generate_charts(self):
        """生成图表"""
        if not HAS_MATPLOTLIB:
            messagebox.showwarning("提示", "需要安装matplotlib库")
            return
        
        if not self.all_notes_data:
            # 从最新文件加载
            self._load_latest_data()
        
        if not self.all_notes_data:
            messagebox.showinfo("提示", "没有数据可分析")
            return
        
        df = pd.DataFrame(self.all_notes_data)
        charts = DataAnalyzer.generate_charts(df, "data/charts")
        
        if charts:
            messagebox.showinfo("完成", f"已生成 {len(charts)} 个图表\n保存到: data/charts/")
            os.startfile("data/charts")
        else:
            messagebox.showwarning("提示", "图表生成失败")
    
    def _generate_wordcloud(self):
        """生成词云"""
        if not HAS_WORDCLOUD:
            messagebox.showwarning("提示", "需要安装wordcloud和jieba库")
            return
        
        if not self.all_notes_data:
            self._load_latest_data()
        
        if not self.all_notes_data:
            messagebox.showinfo("提示", "没有数据可分析")
            return
        
        texts = [d.get('title', '') + ' ' + d.get('content', '') for d in self.all_notes_data]
        output = "data/wordcloud.png"
        
        result = DataAnalyzer.generate_wordcloud(texts, output)
        if result:
            messagebox.showinfo("完成", f"词云已生成: {output}")
            os.startfile(output)
        else:
            messagebox.showwarning("提示", "词云生成失败")
    
    def _generate_report(self):
        """生成分析报告"""
        if not HAS_DOCX:
            messagebox.showwarning("提示", "需要安装python-docx库")
            return
        
        if not self.all_notes_data:
            self._load_latest_data()
        
        if not self.all_notes_data:
            messagebox.showinfo("提示", "没有数据可分析")
            return
        
        df = pd.DataFrame(self.all_notes_data)
        stats = DataAnalyzer.generate_stats(df)
        
        # 先生成图表
        charts = []
        if HAS_MATPLOTLIB:
            charts = DataAnalyzer.generate_charts(df, "data/charts")
        
        keyword = self.all_notes_data[0].get('keyword', '未知') if self.all_notes_data else '未知'
        output = f"data/分析报告_{keyword}_{int(time.time())}.docx"
        
        result = DataAnalyzer.generate_report(df, stats, charts, output, keyword)
        if result:
            messagebox.showinfo("完成", f"报告已生成: {output}")
            os.startfile(output)
        else:
            messagebox.showwarning("提示", "报告生成失败")
    
    def _load_latest_data(self):
        """加载最新数据文件"""
        if not os.path.exists("data"):
            return
        
        files = [f for f in os.listdir("data") if f.startswith("搜索结果_") and f.endswith(".xlsx")]
        if not files:
            return
        
        files.sort(key=lambda x: os.path.getmtime(os.path.join("data", x)), reverse=True)
        latest = os.path.join("data", files[0])
        
        try:
            df = pd.read_excel(latest)
            self.all_notes_data = df.to_dict('records')
        except:
            pass
    
    def _merge_data(self):
        """合并所有数据"""
        if not os.path.exists("data"):
            messagebox.showinfo("提示", "没有数据文件")
            return
        
        all_dfs = []
        for f in os.listdir("data"):
            if f.startswith("搜索结果_") and f.endswith(".xlsx"):
                try:
                    df = pd.read_excel(os.path.join("data", f))
                    all_dfs.append(df)
                except:
                    continue
        
        if not all_dfs:
            messagebox.showinfo("提示", "没有可合并的数据")
            return
        
        merged = pd.concat(all_dfs, ignore_index=True)
        if 'note_link' in merged.columns:
            merged = merged.drop_duplicates(subset=['note_link'])
        
        output = f"data/合并数据_{int(time.time())}.xlsx"
        merged.to_excel(output, index=False)
        
        messagebox.showinfo("完成", f"已合并 {len(merged)} 条数据\n保存到: {output}")
    
    def _refresh_history(self):
        """刷新历史"""
        for item in self.history_tree.get_children():
            self.history_tree.delete(item)
        
        if not os.path.exists("data"):
            return
        
        files = []
        for f in os.listdir("data"):
            if f.startswith("搜索结果_") and f.endswith((".xlsx", ".csv", ".json")):
                path = os.path.join("data", f)
                files.append((f, os.path.getmtime(path), path))
        
        files.sort(key=lambda x: x[1], reverse=True)
        
        for f, mtime, path in files[:20]:
            try:
                keyword = f.replace("搜索结果_", "").rsplit("_", 1)[0]
                time_str = datetime.fromtimestamp(mtime).strftime("%m-%d %H:%M")
                
                if f.endswith(".xlsx"):
                    df = pd.read_excel(path)
                elif f.endswith(".csv"):
                    df = pd.read_csv(path)
                else:
                    df = pd.read_json(path)
                
                notes = len(df)
                images = df['image_count'].sum() if 'image_count' in df.columns else 0
                
                self.history_tree.insert("", tk.END, values=(time_str, keyword, notes, images, f))
            except:
                continue
    
    # === 工具方法 ===
    def _zip_images(self):
        """打包图片"""
        if not os.path.exists("images"):
            messagebox.showinfo("提示", "没有图片目录")
            return
        
        output = f"data/图片打包_{int(time.time())}.zip"
        os.makedirs("data", exist_ok=True)
        
        with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zf:
            for root, dirs, files in os.walk("images"):
                for file in files:
                    filepath = os.path.join(root, file)
                    arcname = os.path.relpath(filepath, "images")
                    zf.write(filepath, arcname)
        
        messagebox.showinfo("完成", f"图片已打包: {output}")
    
    def _open_data_dir(self):
        os.makedirs("data", exist_ok=True)
        os.startfile(os.path.abspath("data"))
    
    def _open_log_file(self):
        if os.path.exists(self.config.log_file):
            os.startfile(self.config.log_file)
        else:
            messagebox.showinfo("提示", "日志文件不存在")
    
    def _clear_log_file(self):
        if os.path.exists(self.config.log_file):
            os.remove(self.config.log_file)
            messagebox.showinfo("完成", "日志已清空")
    
    def _browse_db_path(self):
        path = filedialog.asksaveasfilename(
            defaultextension=".db",
            filetypes=[("SQLite数据库", "*.db")]
        )
        if path:
            self.db_path_var.set(path)
    
    def _on_closing(self):
        """程序退出时的处理"""
        # 保存当前配置
        self._save_gui_settings()
        self.config.save_to_file()
        
        if self.is_running:
            if messagebox.askyesno("确认", "爬取正在进行中，确定要退出吗？"):
                self.should_stop = True
                # 等待一下让爬取线程有机会停止
                self.root.after(500, self._force_close)
            return
        self._force_close()
    
    def _force_close(self):
        """强制关闭程序"""
        # 关闭浏览器
        if self.browser_page:
            try:
                self.browser_page.quit()
            except Exception:
                pass
        # 关闭下载器
        try:
            self.downloader.close()
        except Exception:
            pass
        # 退出程序
        self.root.destroy()
    
    def run(self):
        self.root.mainloop()


if __name__ == '__main__':
    app = CrawlerApp()
    app.run()
